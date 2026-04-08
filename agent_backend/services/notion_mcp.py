# import requests
# import os

# NOTION_API_URL = "https://api.notion.com/v1"
# NOTION_VERSION = "2022-06-28"
# NOTION_TOKEN = "ntn_O37998512432VQZwDeDwERVMgqD2jAE9hBTRoF6Cu81gJG"

# HEADERS = {
#     "Authorization": f"Bearer {NOTION_TOKEN}",
#     "Notion-Version": NOTION_VERSION,
#     "Content-Type": "application/json"
# }

# def notion_list_databases():
#     url = f"{NOTION_API_URL}/search"
#     data = {
#         "filter": {
#             "property": "object",
#             "value": "database"
#         }
#     }
#     response = requests.post(url, headers=HEADERS, json=data)
#     return response.json()


# def notion_list_pages(database_id: str):
#     url = f"{NOTION_API_URL}/databases/{database_id}/query"
#     response = requests.post(url, headers=HEADERS)
#     return response.json()

# def notion_create_page(database_id: str, title: str, content: str = ""):
#     url = f"{NOTION_API_URL}/pages"
#     data = {
#         "parent": {"database_id": database_id},
#         "properties": {
#             "title": [
#                 {
#                     "text": {"content": title}
#                 }
#             ]
#         },
#         "children": [
#             {
#                 "object": "block",
#                 "type": "paragraph",
#                 "paragraph": {
#                     "text": [
#                         {
#                             "type": "text",
#                             "text": {"content": content}
#                         }
#                     ]
#                 }
#             }
#         ] if content else []
#     }
#     response = requests.post(url, headers=HEADERS, json=data)
#     return response.json()

# def notion_search(query: str):
#     url = f"{NOTION_API_URL}/search"
#     data = {"query": query}
#     response = requests.post(url, headers=HEADERS, json=data)
#     return response.json()

# def notion_append_block(page_id: str, content: str):
#     url = f"{NOTION_API_URL}/blocks/{page_id}/children"
#     data = {
#         "children": [
#             {
#                 "object": "block",
#                 "type": "paragraph",
#                 "paragraph": {
#                     "text": [
#                         {
#                             "type": "text",
#                             "text": {"content": content}
#                         }
#                     ]
#                 }
#             }
#         ]
#     }
#     response = requests.patch(url, headers=HEADERS, json=data)
#     return response.json()

# def notion_get_page_content(page_id: str):
#     url = f"{NOTION_API_URL}/blocks/{page_id}/children"
#     response = requests.get(url, headers=HEADERS)
#     return response.json()

# def notion_update_page_title(page_id: str, new_title: str):
#     url = f"{NOTION_API_URL}/pages/{page_id}"
#     data = {
#         "properties": {
#             "title": [
#                 {
#                     "text": {"content": new_title}
#                 }
#             ]
#         }
#     }
#     response = requests.patch(url, headers=HEADERS, json=data)
#     return response.json()

# def notion_delete_page(page_id: str):
#     url = f"{NOTION_API_URL}/blocks/{page_id}"
#     data = {"archived": True}
#     response = requests.patch(url, headers=HEADERS, json=data)
#     return response.json()

# def notion_add_todo(page_id: str, todo_text: str):
#     url = f"{NOTION_API_URL}/blocks/{page_id}/children"
#     data = {
#         "children": [
#             {
#                 "object": "block",
#                 "type": "to_do",
#                 "to_do": {
#                     "text": [
#                         {
#                             "type": "text",
#                             "text": {"content": todo_text}
#                         }
#                     ],
#                     "checked": False
#                 }
#             }
#         ]
#     }
#     response = requests.patch(url, headers=HEADERS, json=data)
#     return response.json()

# def notion_query_database(database_id: str, status_value: str):
#     url = f"{NOTION_API_URL}/databases/{database_id}/query"
#     data = {
#         "filter": {
#             "property": "Status",
#             "select": {
#                 "equals": status_value
#             }
#         }
#     }
#     response = requests.post(url, headers=HEADERS, json=data)
#     return response.json()
# #

from notion_client import Client
import os
from typing import List, Dict, Any, Optional
import json
import io
from docx import Document
import requests
from clients.db_method import get_user_tool_access_token
from bs4 import BeautifulSoup
import math
import uuid
import re
import time

# Notion property types that are read-only / computed and cannot be updated
NON_WRITABLE_PROPERTY_TYPES = {
    "created_time",
    "last_edited_time",
    "created_by",
    "last_edited_by",
    "formula",
    "rollup",
}

# Notion rich_text annotation colors (foreground + background).
NOTION_TEXT_COLORS = frozenset(
    {
        "default",
        "gray",
        "brown",
        "orange",
        "yellow",
        "green",
        "blue",
        "purple",
        "pink",
        "red",
        "gray_background",
        "brown_background",
        "orange_background",
        "yellow_background",
        "green_background",
        "blue_background",
        "purple_background",
        "pink_background",
        "red_background",
    }
)

_NOTION_FG_COLOR_NAMES = frozenset(
    {"gray", "brown", "orange", "yellow", "green", "blue", "purple", "pink", "red"}
)

#
# NOTE: These regexes are used with `.match(text, i)` while scanning a string,
# so they must NOT be anchored to the beginning of the whole string (`^`).
#
_RE_FG_OPEN = re.compile(r"<fg\s+([a-z0-9_]+)\s*>", re.I)
_RE_BG_OPEN = re.compile(r"<bg\s+([a-z0-9_]+)\s*>", re.I)
_RE_FG_CLOSE = re.compile(r"</fg\s*>", re.I)
_RE_BG_CLOSE = re.compile(r"</bg\s*>", re.I)

_RE_B_OPEN = re.compile(r"<\s*(b|strong)\s*>", re.I)
_RE_B_CLOSE = re.compile(r"</\s*(b|strong)\s*>", re.I)


def _notion_color_from_fg_token(token: str) -> str | None:
    t = (token or "").strip().lower()
    if t == "default":
        return "default"
    if t.endswith("_background"):
        return None
    if t in _NOTION_FG_COLOR_NAMES:
        return t
    return None


def _notion_color_from_bg_token(token: str) -> str | None:
    t = (token or "").strip().lower()
    if t.endswith("_background") and t in NOTION_TEXT_COLORS:
        return t
    if t in _NOTION_FG_COLOR_NAMES:
        return f"{t}_background"
    return None


def _rich_text_from_markup(text: str) -> List[Dict[str, Any]]:
    """
    Convert a simple markdown-like string into Notion rich_text spans
    with basic formatting:
    - **bold**
    - *italic* (single `*` toggles italic; `**` is bold)
    - __underline__
    - ~~strikethrough~~
    - [label](url) for hyperlinks (Notion native links)
    - <fg COLOR>text</fg> for foreground color (COLOR: gray, brown, orange,
      yellow, green, blue, purple, pink, red, or default)
    - <bg COLOR>text</bg> for highlight / background color (COLOR as above, or
      already suffixed with _background)

    If no formatting markers are detected, returns a single plain span.
    This is intentionally conservative and does not attempt full markdown.
    """
    if not text:
        return []

    # Quick check to avoid regex work when unnecessary
    if not any(
        m in text
        for m in [
            "**",
            "*",
            "__",
            "_",
            "~~",
            "](",
            "<fg",
            "<bg",
            "</fg",
            "</bg",
            "<a ",
            "</a>",
            "<b",
            "</b",
            "<strong",
            "</strong",
        ]
    ):
        return [{"type": "text", "text": {"content": text}}]

    spans: List[Dict[str, Any]] = []
    i = 0
    n = len(text)
    active = {
        "bold": False,
        "italic": False,
        "underline": False,
        "strikethrough": False,
    }
    # Stack of (kind, notion_color); innermost scope wins for make_span.
    color_stack: list[tuple[str, str]] = []

    def current_color() -> str:
        return color_stack[-1][1] if color_stack else "default"

    def make_span(content: str, link_url: str | None = None) -> Dict[str, Any]:
        item: Dict[str, Any] = {"type": "text", "text": {"content": content}}
        if link_url:
            item["text"]["link"] = {"url": link_url}
        color_val = current_color()
        need_ann = any(active.values()) or color_val != "default"
        if need_ann:
            item["annotations"] = {
                "bold": active["bold"],
                "italic": active["italic"],
                "underline": active["underline"],
                "strikethrough": active["strikethrough"],
                "code": False,
                "color": color_val,
            }
        return item

    while i < n:
        # HTML bold tags: <b>...</b> / <strong>...</strong>
        m_b_open = _RE_B_OPEN.match(text, i)
        if m_b_open:
            active["bold"] = True
            i = m_b_open.end()
            continue
        m_b_close = _RE_B_CLOSE.match(text, i)
        if m_b_close:
            active["bold"] = False
            i = m_b_close.end()
            continue

        # HTML anchor support:
        # Convert <a href="https://...">label</a> into a native Notion link span.
        if text.startswith("<a", i):
            # Keep it intentionally simple: url is in href="...", label is direct text.
            m_a = re.match(
                r'^<a\s+href=["\']([^"\']+)["\']\s*>(.*?)</a>',
                text[i:],
                flags=re.IGNORECASE,
            )
            if m_a:
                url = m_a.group(1)
                label = m_a.group(2)
                # Only emit a link if both url and label exist.
                if url and label is not None:
                    spans.append(make_span(label, link_url=url))
                    i += m_a.end()
                    continue

        # Color / highlight tags (must run before '[' so links can be colored)
        m_fg_o = _RE_FG_OPEN.match(text, i)
        if m_fg_o:
            col = _notion_color_from_fg_token(m_fg_o.group(1))
            if col:
                color_stack.append(("fg", col))
                i = m_fg_o.end()
                continue
            spans.append(make_span(text[i]))
            i += 1
            continue
        m_bg_o = _RE_BG_OPEN.match(text, i)
        if m_bg_o:
            col = _notion_color_from_bg_token(m_bg_o.group(1))
            if col:
                color_stack.append(("bg", col))
                i = m_bg_o.end()
                continue
            spans.append(make_span(text[i]))
            i += 1
            continue
        m_fg_c = _RE_FG_CLOSE.match(text, i)
        if m_fg_c:
            if color_stack and color_stack[-1][0] == "fg":
                color_stack.pop()
                i = m_fg_c.end()
                continue
            spans.append(make_span(text[i]))
            i += 1
            continue
        m_bg_c = _RE_BG_CLOSE.match(text, i)
        if m_bg_c:
            if color_stack and color_stack[-1][0] == "bg":
                color_stack.pop()
                i = m_bg_c.end()
                continue
            spans.append(make_span(text[i]))
            i += 1
            continue

        # Link [label](url)
        if text[i] == "[":
            close = text.find("]", i + 1)
            if close != -1 and close + 1 < n and text[close + 1] == "(":
                end_paren = text.find(")", close + 2)
                if end_paren != -1:
                    label = text[i + 1 : close]
                    url = text[close + 2 : end_paren]
                    if label and url:
                        spans.append(make_span(label, link_url=url))
                        i = end_paren + 1
                        continue

        # Toggle style markers (stackable)
        if text.startswith("**", i):
            active["bold"] = not active["bold"]
            i += 2
            continue
        if text.startswith("__", i):
            active["underline"] = not active["underline"]
            i += 2
            continue
        if text.startswith("~~", i):
            active["strikethrough"] = not active["strikethrough"]
            i += 2
            continue
        if text[i] == "*":
            prev_ch = text[i - 1] if i > 0 else ""
            next_ch = text[i + 1] if i + 1 < n else ""
            # Avoid treating '*' inside words as italic markers (e.g. "foo*bar").
            if prev_ch.isalnum() and next_ch.isalnum():
                spans.append(make_span("*"))
                i += 1
                continue
            active["italic"] = not active["italic"]
            i += 1
            continue
        if text[i] == "_":
            # Avoid treating '_' inside words as italic markers (e.g. "is_palindrome").
            prev_ch = text[i - 1] if i > 0 else ""
            next_ch = text[i + 1] if i + 1 < n else ""
            if prev_ch.isalnum() and next_ch.isalnum():
                spans.append(make_span("_"))
                i += 1
                continue
            # Single underscore toggles italic; double underscore is handled above as underline.
            active["italic"] = not active["italic"]
            i += 1
            continue

        # Collect plain run until next marker or link opener
        start = i
        while (
            i < n
            and not text.startswith("**", i)
            and not text.startswith("__", i)
            and not text.startswith("~~", i)
            and not text.startswith("<fg", i)
            and not text.startswith("</fg", i)
            and not text.startswith("<bg", i)
            and not text.startswith("</bg", i)
            and text[i] != "*"
            and text[i] != "_"
            and text[i] != "["
            and not text.startswith("<a", i)
            and not text.startswith("<b", i)
            and not text.startswith("</b", i)
            and not text.startswith("<strong", i)
            and not text.startswith("</strong", i)
        ):
            i += 1

        if i > start:
            spans.append(make_span(text[start:i]))
        else:
            # Safety advance to avoid non-progress loops on odd characters.
            spans.append(make_span(text[i]))
            i += 1

    return spans


def _content_has_rich_markup(text: str) -> bool:
    """True if text may contain tokens handled by _rich_text_from_markup."""
    if not text:
        return False
    return any(
        m in text
        for m in (
            "**",
            "__",
            "~~",
            "](",
            "<fg",
            "<bg",
            "</fg",
            "</bg",
            "<a ",
            "</a>",
            "<b",
            "</b",
            "<strong",
            "</strong",
        )
    )


def _merge_uniform_into_rich_text_spans(
    spans: List[Dict[str, Any]],
    bold: bool | None,
    italic: bool | None,
    underline: bool | None,
    strikethrough: bool | None,
    text_color: str | None,
) -> None:
    """
    Apply tool-level uniform color / emphasis onto parsed rich_text spans.
    Only True flags and valid text_color are applied so explicit False from
    clients (common in JSON tool calls) does not disable markdown parsing.
    """
    ucolor: str | None = None
    if text_color is not None and str(text_color).strip():
        raw = str(text_color).strip().lower()
        if raw in NOTION_TEXT_COLORS:
            ucolor = raw
    if not (
        ucolor is not None
        or bold is True
        or italic is True
        or underline is True
        or strikethrough is True
    ):
        return
    for item in spans:
        if not isinstance(item, dict) or item.get("type") != "text":
            continue
        ann = item.get("annotations")
        if ann is None:
            ann = {
                "bold": False,
                "italic": False,
                "underline": False,
                "strikethrough": False,
                "code": False,
                "color": "default",
            }
            item["annotations"] = ann
        if ucolor is not None and ann.get("color", "default") == "default":
            ann["color"] = ucolor
        if bold is True:
            ann["bold"] = True
        if italic is True:
            ann["italic"] = True
        if underline is True:
            ann["underline"] = True
        if strikethrough is True:
            ann["strikethrough"] = True


def _split_text_for_notion_limit(text: str, max_len: int = 2000) -> List[str]:
    """Split long text into <= max_len chunks, preferring newline boundaries."""
    if len(text) <= max_len:
        return [text]
    chunks: List[str] = []
    remaining = text
    while remaining:
        if len(remaining) <= max_len:
            chunks.append(remaining)
            break
        cut = remaining.rfind("\n", 0, max_len + 1)
        if cut <= 0:
            cut = max_len
        chunks.append(remaining[:cut])
        remaining = remaining[cut:].lstrip("\n")
    return chunks


def _blocks_from_text_content(content: str) -> List[Dict[str, Any]]:
    """
    Convert markdown-like multiline text into Notion blocks.
    Supports:
    - Headings: '#', '##', '###'
    - Fenced code blocks: ```lang ... ```
    - Unordered lists: '- item', '* item'
    - Ordered lists: '1. item', '2. item'
    - Fallback paragraph lines
    Inline formatting in each line is parsed via _rich_text_from_markup.
    """
    content = (content or "").replace("\r\n", "\n")
    if not content.strip():
        return []

    blocks: List[Dict[str, Any]] = []
    ul_re = re.compile(r"^([-*•])\s+(.*)$")
    ol_re = re.compile(r"^(\d+)[\.\)]\s+(.*)$")
    heading_re = re.compile(r"^(#{1,3})\s+(.*)$")
    code_fence_re = re.compile(r"^```([\w+-]*)\s*$")

    list_types = {"bulleted_list_item", "numbered_list_item"}

    def _append_child(parent: Dict[str, Any], child: Dict[str, Any]) -> None:
        ptype = parent.get("type")
        if not ptype:
            return
        parent.setdefault(ptype, {})
        parent[ptype].setdefault("children", [])
        parent[ptype]["children"].append(child)

    def _last_top_level_list_item() -> Dict[str, Any] | None:
        for b in reversed(blocks):
            if b.get("type") in list_types:
                return b
        return None

    lines = content.split("\n")
    in_code_block = False
    code_language = "plain text"
    code_lines: List[str] = []

    def _flush_code_block() -> None:
        nonlocal code_lines, code_language
        if not code_lines:
            return
        code_text = "\n".join(code_lines)
        for segment in _split_text_for_notion_limit(code_text):
            blocks.append(
                {
                    "object": "block",
                    "type": "code",
                    "code": {
                        "rich_text": [{"type": "text", "text": {"content": segment}}],
                        "language": code_language or "plain text",
                    },
                }
            )
        code_lines = []
        code_language = "plain text"

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        fence = code_fence_re.match(line.strip())
        if fence:
            if in_code_block:
                _flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
                lang = (fence.group(1) or "").strip().lower()
                code_language = lang if lang else "plain text"
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if not line.strip():
            continue
        leading_spaces = len(line) - len(line.lstrip(" "))
        stripped = line.lstrip(" ")

        bullet_marker = None
        h = heading_re.match(stripped)
        if h:
            hashes = h.group(1)
            heading_text = h.group(2).strip()
            heading_type = {
                1: "heading_1",
                2: "heading_2",
                3: "heading_3",
            }.get(len(hashes), "heading_3")
            for segment in _split_text_for_notion_limit(heading_text):
                blocks.append(
                    {
                        "object": "block",
                        "type": heading_type,
                        heading_type: {"rich_text": _rich_text_from_markup(segment)},
                    }
                )
            continue

        m_ul = ul_re.match(stripped)
        if m_ul:
            bullet_marker = m_ul.group(1)
            line_type = "bulleted_list_item"
            line_text = m_ul.group(2).strip()
        else:
            m_ol = ol_re.match(stripped)
            if m_ol:
                line_type = "numbered_list_item"
                line_text = m_ol.group(2).strip()
            else:
                line_type = "paragraph"
                line_text = stripped.strip()

        for segment in _split_text_for_notion_limit(line_text):
            new_block = {
                "object": "block",
                "type": line_type,
                line_type: {"rich_text": _rich_text_from_markup(segment)},
            }

            # Nest sub-bullets under the nearest numbered item so OL numbering
            # doesn't reset when explanatory bullet lines are present.
            if line_type == "bulleted_list_item":
                attach_to_parent = False
                if leading_spaces > 0 and blocks and blocks[-1].get("type") in list_types:
                    attach_to_parent = True
                    parent = blocks[-1]
                elif bullet_marker in {"•", "-", "*"}:
                    # Heuristic for LLM/user text that uses unindented subpoints
                    # under numbered items. Attach to nearest numbered top-level
                    # item to preserve OL numbering continuity.
                    parent = _last_top_level_list_item()
                    attach_to_parent = parent is not None and parent.get("type") == "numbered_list_item"

                if attach_to_parent:
                    _append_child(parent, new_block)
                    continue

            blocks.append(new_block)

    # Unclosed code fence: still emit captured code to avoid data loss.
    if in_code_block:
        _flush_code_block()

    return blocks

# Helper to get a Notion client for a user

# BASE_URL = "http://3.6.95.164:5000/users"

# Example: Call get_tool_token endpoint
# def get_tool_token(unified_token):
#     tool_name = "Notion"
#     url = f"{BASE_URL}/get_tool_token"
#     payload = {
#         "unified_token": unified_token,
#         "tool_name": tool_name
#     }

#     response = requests.post(url, json=payload)
#     print("I am here in get_tool_token", response)
#     print("abc")
#     if response.status_code == 200:
#         print("Access Token:", response.json())
#         return response.json()
#     else:
#         print("Error:", response.status_code, response.text)


def get_notion_client(unified_token):
    # result = get_tool_token(unified_token)
    tool_name = "Notion"
    result, status = get_user_tool_access_token(unified_token, tool_name)
    
    # Check if credentials exist before accessing
    if status != 200 or not isinstance(result, dict) or "access_token" not in result:
        error_msg = result.get("error", "Failed to retrieve Notion credentials") if isinstance(result, dict) else "Failed to retrieve Notion credentials"
        raise Exception(f"Failed to retrieve Notion credentials. Please connect Notion. {error_msg}")
    
    access_token = result.get("access_token", "")
    # Handle both {"secret": "..."} and direct string
    if isinstance(access_token, dict) and "secret" in access_token:
        notion_token = access_token["secret"]
    elif isinstance(access_token, str):
        notion_token = access_token
    else:
        raise Exception("Notion access token format invalid.")
    return Client(auth=notion_token)


def get_title_from_properties(properties: Dict[str, Any]) -> str:
    """
    Safely extract a title string from a page/database properties object.
    Falls back to \"(no title)\" if no title property is present or non-empty.
    """
    for prop in properties.values():
        if prop.get("type") == "title":
            pieces: list[str] = []
            for part in prop.get("title", []):
                # Newer Notion responses expose plain_text; older ones expose text.content
                if "plain_text" in part:
                    pieces.append(part.get("plain_text") or "")
                elif "text" in part:
                    pieces.append(part["text"].get("content", ""))
            title = "".join(pieces).strip()
            return title or "(no title)"
    return "(no title)"


def extract_property_value(prop_obj: Dict[str, Any]) -> Any:
    """
    Generic extractor that turns a Notion page property object into a simple Python value.
    Supports most core property types (title, rich_text, select, status, multi_select,
    people, relation, number, checkbox, date, url, email, phone_number, files,
    formula, rollup, created_* and last_edited_*).
    """
    if not isinstance(prop_obj, dict):
        return None

    prop_type = prop_obj.get("type")

    if prop_type == "title":
        pieces: list[str] = []
        for part in prop_obj.get("title", []):
            if "plain_text" in part:
                pieces.append(part.get("plain_text") or "")
            elif "text" in part:
                pieces.append(part["text"].get("content", ""))
        return "".join(pieces).strip() or None

    if prop_type == "rich_text":
        parts: list[str] = []
        for rt in prop_obj.get("rich_text", []):
            if "plain_text" in rt:
                parts.append(rt.get("plain_text") or "")
            elif "text" in rt:
                parts.append(rt["text"].get("content", ""))
        text = " ".join(p.strip() for p in parts if p.strip())
        return text or None

    if prop_type == "number":
        return prop_obj.get("number")

    if prop_type == "checkbox":
        return prop_obj.get("checkbox")

    if prop_type in ("select", "status"):
        value = prop_obj.get(prop_type)
        if isinstance(value, dict):
            return value.get("name")
        return None

    if prop_type == "multi_select":
        return [
            opt.get("name")
            for opt in prop_obj.get("multi_select", [])
            if isinstance(opt, dict) and opt.get("name") is not None
        ]

    if prop_type == "relation":
        return [
            rel.get("id")
            for rel in prop_obj.get("relation", [])
            if isinstance(rel, dict) and rel.get("id")
        ]

    if prop_type == "url":
        return prop_obj.get("url")

    if prop_type == "email":
        return prop_obj.get("email")

    if prop_type == "phone_number":
        return prop_obj.get("phone_number")

    if prop_type == "date":
        date_obj = prop_obj.get("date")
        if isinstance(date_obj, dict):
            # Return the start ISO timestamp; callers can inspect end/range if needed.
            return date_obj.get("start")
        return None

    if prop_type == "people":
        people: list[str] = []
        for p in prop_obj.get("people", []):
            if not isinstance(p, dict):
                continue
            name = p.get("name")
            email = None
            person_info = p.get("person")
            if isinstance(person_info, dict):
                email = person_info.get("email")
            people.append(name or email or p.get("id"))
        return people

    if prop_type == "files":
        files: list[Dict[str, Any]] = []
        for f in prop_obj.get("files", []):
            if not isinstance(f, dict):
                continue
            name = f.get("name")
            file_type = f.get("type")
            url = None
            if file_type == "external":
                url = f.get("external", {}).get("url")
            elif file_type == "file":
                url = f.get("file", {}).get("url")
            files.append({"name": name, "url": url, "type": file_type})
        return files

    if prop_type == "formula":
        formula = prop_obj.get("formula", {})
        if isinstance(formula, dict):
            ftype = formula.get("type")
            return formula.get(ftype)
        return None

    if prop_type == "rollup":
        rollup = prop_obj.get("rollup", {})
        if not isinstance(rollup, dict):
            return None
        rtype = rollup.get("type")
        if rtype == "array":
            array_values = []
            for item in rollup.get("array", []):
                if isinstance(item, dict):
                    # Items are themselves property-like objects
                    array_values.append(extract_property_value(item))
            return array_values
        return rollup.get(rtype)

    if prop_type in ("created_time", "last_edited_time"):
        return prop_obj.get(prop_type)

    if prop_type in ("created_by", "last_edited_by"):
        user = prop_obj.get(prop_type)
        if isinstance(user, dict):
            return user.get("name") or user.get("id")
        return None

    # Fallback: return the raw object so callers can inspect if needed
    return None


# Simple in-memory cache for Notion user lookups (email -> user_id)
_USER_EMAIL_CACHE: Dict[str, str] = {}


def _with_notion_rate_limit_retry(func, *args, **kwargs):
    """
    Execute a Notion API call with basic retry/backoff on rate limiting.
    We intentionally keep this very lightweight and avoid depending on SDK error types.
    """
    max_attempts = 3
    backoff = 0.5
    for attempt in range(max_attempts):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            msg = str(e).lower()
            if (
                "rate_limited" in msg
                or "too many requests" in msg
                or "status code: 429" in msg
                or "http 429" in msg
            ) and attempt < max_attempts - 1:
                time.sleep(backoff)
                backoff *= 2
                continue
            raise


def notion_list_databases(unified_token: str = None) -> List[Dict]:
    notion = get_notion_client(unified_token)

    databases: List[Dict[str, Any]] = []
    start_cursor = None

    while True:
        response = _with_notion_rate_limit_retry(
            notion.search,
            filter={"property": "object", "value": "database"},
            start_cursor=start_cursor,
        )
        results = response.get("results", [])

        for db in results:
            title_parts = db.get("title", []) or []
            pieces: list[str] = []
            for part in title_parts:
                if "plain_text" in part:
                    pieces.append(part.get("plain_text") or "")
                elif "text" in part:
                    pieces.append(part["text"].get("content", ""))
            title = "".join(pieces).strip() or "(no title)"
            databases.append(
                {
                    "id": db.get("id"),
                    "title": title,
                    "url": db.get("url"),
                }
            )

        if not response.get("has_more"):
            break
        start_cursor = response.get("next_cursor")

    return databases


def notion_list_pages(database_id: str, unified_token: str = None) -> List[Dict]:
    notion = get_notion_client(unified_token)
    
    # First check if the provided ID is a database or page
    try:
        # Try to retrieve as database first
        db_info = _with_notion_rate_limit_retry(notion.databases.retrieve, database_id)
        # If successful, it's a database - proceed with querying pages
        pages = []
        has_more = True
        start_cursor = None

        while has_more:
            query = {"database_id": database_id}
            if start_cursor:
                query["start_cursor"] = start_cursor

            response = _with_notion_rate_limit_retry(notion.databases.query, **query)
            results = response.get("results", [])

            for p in results:
                title = get_title_from_properties(p.get("properties", {}))
                pages.append({"id": p["id"], "title": title})

            has_more = response.get("has_more", False)
            start_cursor = response.get("next_cursor")

        return pages
    except Exception as e:
        # If database retrieval fails, check if it's a page
        try:
            page_info = notion.pages.retrieve(database_id)
            # It's a page, not a database - return child pages/databases
            return notion_list_child_pages(database_id, unified_token)
        except:
            # Neither page nor database
            raise ValueError(f"Invalid ID {database_id}. Please provide a valid database or page ID")


def notion_find_or_create_database(database_name: str, unified_token: str = None) -> Dict:
    """Find a database by name, or suggest creating one if not found"""
    notion = get_notion_client(unified_token)
    
    try:
        # Search for databases with the given name
        results = notion.search(
            query=database_name,
            filter={"property": "object", "value": "database"}
        )["results"]
        
        # Look for exact or partial matches
        exact_matches = []
        partial_matches = []
        
        for db in results:
            db_title = ""
            if "title" in db and db["title"]:
                db_title = db["title"][0]["plain_text"] if db["title"][0].get("plain_text") else ""
            
            if db_title.lower() == database_name.lower():
                exact_matches.append({
                    "id": db["id"],
                    "title": db_title,
                    "url": db.get("url", ""),
                    "match_type": "exact"
                })
            elif database_name.lower() in db_title.lower():
                partial_matches.append({
                    "id": db["id"],
                    "title": db_title,
                    "url": db.get("url", ""),
                    "match_type": "partial"
                })
        
        if exact_matches:
            return {
                "success": True,
                "found": True,
                "databases": exact_matches,
                "message": f"Found {len(exact_matches)} exact match(es) for '{database_name}'"
            }
        elif partial_matches:
            return {
                "success": True,
                "found": True,
                "databases": partial_matches,
                "message": f"Found {len(partial_matches)} partial match(es) for '{database_name}'"
            }
        else:
            # Get available pages (including nested pages) for creating a new database
            parent_pages = notion_list_all_database_parents(unified_token)
            return {
                "success": True,
                "found": False,
                "databases": [],
                "available_parents": parent_pages[:5],  # Show first 5 parent-capable pages
                "message": f"No database named '{database_name}' found. You can create one using create_database function.",
                "suggestion": f"Use create_database with a parent_id (any page, including nested pages) to create '{database_name}' database"
            }
            
    except Exception as e:
        return {"success": False, "error": f"Error searching for database: {str(e)}"}


def notion_create_page(
    parent_id: str = None, 
    title: str = "Untitled", 
    unified_token: str = None, 
    content: str = "",
    database_id: str = None,  # Added for backward compatibility
    database_name: str = None  # New: find database by name
) -> Dict:
    """
    Create a new Notion page.
    
    Args:
        parent_id: The ID of the parent page or database (legacy parameter)
        title: The title of the new page
        unified_token: Authentication token
        content: Optional content for the page
        database_id: The ID of the database to create the page in (alternative to parent_id)
        database_name: Name of database to find and create page in (alternative to IDs)
    """
    notion = get_notion_client(unified_token)
    
    # If database_name is provided, try to find the database first
    if database_name and not database_id and not parent_id:
        db_search = notion_find_or_create_database(database_name, unified_token)
        if db_search["success"] and db_search["found"] and db_search["databases"]:
            database_id = db_search["databases"][0]["id"]  # Use first match
        else:
            return {
                "success": False, 
                "error": f"Database '{database_name}' not found. {db_search.get('suggestion', 'Please create the database first.')}",
                "available_parents": db_search.get("available_parents", []),
                "search_result": db_search
            }
    
    # Handle both database_id and parent_id for backward compatibility
    target_id = database_id or parent_id
    if not target_id:
        # Get available options for user
        databases = notion_list_databases(unified_token)
        parent_pages = notion_list_parent_pages(unified_token)
        
        return {
            "success": False, 
            "error": "Either parent_id, database_id, or database_name must be provided",
            "available_databases": databases[:5] if databases else [],
            "available_parent_pages": parent_pages[:5] if parent_pages else [],
            "suggestion": "Use list_databases() to see available databases, or create_database() to create a new one"
        }
    
    try:
        # First try to treat the target as a database
        try:
            db = notion.databases.retrieve(target_id)
            title_prop_name = next(
                (k for k, v in db["properties"].items() if v["type"] == "title"), "Name"
            )
            
            # Create page in database (always set title; content handled as a paragraph)
            page = notion.pages.create(
                parent={"database_id": target_id},
                properties={title_prop_name: {"title": [{"text": {"content": title}}]}},
                children=_blocks_from_text_content(content) if content else [],
            )
            return {"success": True, "id": page["id"], "url": page["url"], "type": "page_in_database", "database_title": db.get("title", [{}])[0].get("plain_text", "Unknown")}
            
        except Exception as db_error:
            # If that fails, try to treat it as a page
            try:
                page_info = notion.pages.retrieve(target_id)
                
                # Create page as child of another page
                page = notion.pages.create(
                    parent={"page_id": target_id},
                    properties={
                        "title": {"title": [{"text": {"content": title}}]}
                    },
                    children=_blocks_from_text_content(content) if content else [],
                )
                return {"success": True, "id": page["id"], "url": page["url"], "type": "child_page"}
                
            except Exception as page_error:
                return {
                    "success": False, 
                    "error": f"Invalid ID '{target_id}'. Not a valid database or page ID.",
                    "details": f"Database error: {str(db_error)}, Page error: {str(page_error)}"
                }
                    
    except Exception as e:
        return {"success": False, "error": f"Failed to create page: {str(e)}"}


def notion_search_notion(query: str, unified_token: str = None) -> List[Dict]:
    notion = get_notion_client(unified_token)
    results = notion.search(query=query)["results"]
    items = []
    for r in results:
        if r["object"] == "page":
            title = get_title_from_properties(r.get("properties", {}))
            items.append({"id": r["id"], "title": title, "url": r.get("url"), "type": "page"})
        elif r["object"] == "database":
            title_parts = r.get("title", [])
            title = "".join(part["text"]["content"] for part in title_parts if "text" in part) or "(no title)"
            items.append({"id": r["id"], "title": title, "url": r.get("url"), "type": "database"})
    return items


def notion_append_block(
    page_id: str,
    content: str | None = None,
    unified_token: str = None,
    mode: str | None = None,
    table: Dict[str, Any] | None = None,
    heading_level: int | None = None,
    after_block_id: str | None = None,
    parent_block_id: str | None = None,
) -> Dict:
    """
    Append new content to a Notion page.

    - If ``table`` is provided, a Notion table block is created using the
      structure in ``table`` (columns + rows) and appended to the page.
    - Otherwise, a paragraph text block is created using ``content``.
    - If ``heading_level`` is 1, 2, or 3, each non-empty line becomes a native
      Notion heading (heading_1 / heading_2 / heading_3) instead of a paragraph.
      Use this when the user asks for a heading without markdown ``#`` prefixes.
    - ``after_block_id``: insert new blocks immediately after this block; it must
      be a direct child of the append parent. From ``get_page_content``, use the
      ``block_id`` of the sibling *before* the insertion point.
    - ``parent_block_id``: append under this block instead of under ``page_id``
      (nested toggles, columns, list items). Omit for page-level children.

    Long text content is split into multiple paragraph blocks to respect
    Notion's 2000-character limit per text block.
    """
    notion = get_notion_client(unified_token)

    append_parent_id = (
        str(parent_block_id).strip()
        if parent_block_id and str(parent_block_id).strip()
        else page_id
    )
    after_kw: Dict[str, Any] = {}
    if after_block_id and str(after_block_id).strip():
        after_kw["after"] = str(after_block_id).strip()

    # Notion has a 2000 character limit per text block
    MAX_CONTENT_LENGTH = 2000

    if mode and mode.lower() == "replace":
        # Fetch and delete all existing children with pagination
        start_cursor = None
        while True:
            response = _with_notion_rate_limit_retry(
                notion.blocks.children.list, block_id=page_id, start_cursor=start_cursor
            )
            existing_children = response.get("results", [])
            for block in existing_children:
                _with_notion_rate_limit_retry(notion.blocks.delete, block_id=block["id"])
            if not response.get("has_more"):
                break
            start_cursor = response.get("next_cursor")
    
    # If a structured table is provided, create a Notion table block.
    if table is not None:
        try:
            columns = table.get("columns") or []
            rows = table.get("rows") or []
            has_column_header = bool(table.get("has_column_header", True))
            has_row_header = bool(table.get("has_row_header", False))

            if not isinstance(columns, list) or not all(
                isinstance(c, str) for c in columns
            ):
                raise ValueError("table.columns must be a list of strings")
            if not isinstance(rows, list) or not all(
                isinstance(r, (list, tuple)) for r in rows
            ):
                raise ValueError(
                    "table.rows must be a list of row arrays (each row is a list of cell strings)"
                )

            table_width = len(columns)
            if table_width == 0:
                raise ValueError("table.columns must contain at least one column")

            def _cell(value: str) -> Dict[str, Any]:
                return {
                    "type": "text",
                    "text": {"content": value if value is not None else ""},
                }

            children: list[Dict[str, Any]] = []

            # Optional header row from columns
            if has_column_header:
                header_cells: list[list[Dict[str, Any]]] = []
                for col in columns:
                    header_cells.append([_cell(str(col))])
                children.append(
                    {
                        "object": "block",
                        "type": "table_row",
                        "table_row": {"cells": header_cells},
                    }
                )

            # Data rows
            for row in rows:
                # Pad or trim row to match table_width
                row_values = list(row)[:table_width]
                if len(row_values) < table_width:
                    row_values.extend([""] * (table_width - len(row_values)))
                row_cells: list[list[Dict[str, Any]]] = []
                for cell_value in row_values:
                    row_cells.append([_cell(str(cell_value))])
                children.append(
                    {
                        "object": "block",
                        "type": "table_row",
                        "table_row": {"cells": row_cells},
                    }
                )

            table_block: Dict[str, Any] = {
                "object": "block",
                "type": "table",
                "table": {
                    "table_width": table_width,
                    "has_column_header": has_column_header,
                    "has_row_header": has_row_header,
                    "children": children,
                },
            }

            _with_notion_rate_limit_retry(
                notion.blocks.children.append,
                block_id=append_parent_id,
                children=[table_block],
                **after_kw,
            )
            return {
                "status": "success",
                "message": "Table appended.",
                "table_width": table_width,
                "row_count": len(children),
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to append table: {str(e)}",
            }

    # Fallback: text content. None/empty becomes empty string to keep behaviour
    # backward compatible for callers that always passed a string.
    if content is None:
        content = ""

    heading_level_int: int | None = None
    if heading_level is not None:
        try:
            h = int(heading_level)
            if h in (1, 2, 3):
                heading_level_int = h
        except (TypeError, ValueError):
            heading_level_int = None

    if heading_level_int is not None and content.strip():
        heading_type = f"heading_{heading_level_int}"
        heading_blocks: List[Dict[str, Any]] = []
        for raw_line in content.replace("\r\n", "\n").split("\n"):
            line = raw_line.strip()
            if not line:
                continue
            for segment in _split_text_for_notion_limit(line):
                heading_blocks.append(
                    {
                        "object": "block",
                        "type": heading_type,
                        heading_type: {
                            "rich_text": _rich_text_from_markup(segment),
                        },
                    }
                )
        if not heading_blocks:
            return {"status": "success", "message": "No content to append."}
        _with_notion_rate_limit_retry(
            notion.blocks.children.append,
            block_id=append_parent_id,
            children=heading_blocks,
            **after_kw,
        )
        return {
            "status": "success",
            "message": "Content appended.",
            "blocks_created": len(heading_blocks),
        }

    blocks = _blocks_from_text_content(content)
    if not blocks:
        return {"status": "success", "message": "No content to append."}

    # Smart nesting for bullet-only appends.
    # 1) If bullet count matches remaining numbered items without children,
    #    distribute one bullet under each numbered item in order.
    # 2) Otherwise, attach all bullets to the next numbered item without
    #    children (or the last numbered item as fallback).
    only_bullets = all(b.get("type") == "bulleted_list_item" for b in blocks)
    appended_to_target = False
    # Explicit position: skip smart list nesting so after_kw applies.
    if only_bullets and not after_kw:
        try:
            children_resp = _with_notion_rate_limit_retry(
                notion.blocks.children.list, block_id=append_parent_id
            )
            top_blocks = children_resp.get("results", [])
            numbered_items = [
                b for b in top_blocks if b.get("type") == "numbered_list_item" and b.get("id")
            ]
            def _item_has_children(item: Dict[str, Any]) -> bool:
                # Notion commonly exposes child presence via top-level
                # `has_children`; keep payload fallback for completeness.
                if bool(item.get("has_children")):
                    return True
                payload = item.get("numbered_list_item", {}) or {}
                return bool(payload.get("children"))
            available_numbered = [
                item
                for item in numbered_items
                if not _item_has_children(item)
            ]

            # Distribute bullets one-by-one when counts align
            if available_numbered and len(blocks) == len(available_numbered):
                for idx, bullet_block in enumerate(blocks):
                    _with_notion_rate_limit_retry(
                        notion.blocks.children.append,
                        block_id=available_numbered[idx]["id"],
                        children=[bullet_block],
                    )
                appended_to_target = True
            else:
                target_parent = available_numbered[0] if available_numbered else None
                if not target_parent and numbered_items:
                    target_parent = numbered_items[-1]
                if target_parent:
                    _with_notion_rate_limit_retry(
                        notion.blocks.children.append,
                        block_id=target_parent["id"],
                        children=blocks,
                    )
                    appended_to_target = True
        except Exception:
            appended_to_target = False

    if not appended_to_target:
        _with_notion_rate_limit_retry(
            notion.blocks.children.append,
            block_id=append_parent_id,
            children=blocks,
            **after_kw,
        )
    return {
        "status": "success",
        "message": "Content appended.",
        "blocks_created": len(blocks),
    }


def _parse_table_content_to_rows(content: str, table_width: int) -> List[List[str]]:
    """
    Parse a string into rows of cell values for appending to a Notion table.
    Supports: markdown-style (| a | b |), pipe-separated, tab-separated;
    one row per line or a single line for one row.
    """
    content = (content or "").strip()
    if not content:
        return []
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]

    # Markdown-style: lines with | ... |
    if "|" in content and content.count("|") >= 2:
        rows = []
        for line in lines:
            # Skip separator lines like |---|-----|
            if re.match(r"^[\s|\-]+$", line.replace(" ", "")):
                continue
            parts = line.split("|")
            # | a | b | gives ['', ' a ', ' b ', ''] -> take [1:-1] and strip
            cells = [p.strip() for p in parts[1:-1]] if len(parts) > 2 else [p.strip() for p in parts if p.strip()]
            if cells:
                rows.append(cells)
        return rows

    # Tab-separated (one or more lines)
    if "\t" in content:
        return [[c.strip() for c in line.split("\t")] for line in lines]

    # Pipe-separated (space-pipe-space)
    if " | " in content:
        if "\n" not in content:
            return [[c.strip() for c in content.split(" | ")]]
        return [[c.strip() for c in line.split(" | ")] for line in lines]

    # Single line, no delimiter: one cell
    if len(lines) == 1:
        return [[content]]
    return [[ln] for ln in lines]


def notion_update_block(
    block_id: str,
    content: str | None,
    unified_token: str = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    strikethrough: bool | None = None,
    text_color: str | None = None,
) -> Dict:
    """
    Update/edit a specific block in a Notion page by its block ID.
    This is better than delete+append as it preserves the block in place.
    Handles content longer than 2000 characters by splitting into multiple blocks.
    For table blocks: appends new row(s) from parsed content (markdown, tab-, or pipe-separated).
    For table_row blocks: replaces the row cells with parsed content.
    """
    notion = get_notion_client(unified_token)

    def _append_table_rows(block_id: str, content: str, table_width: int) -> Dict:
        """Append parsed row(s) to a table block (used when retrieve fails or block is table)."""
        rows = _parse_table_content_to_rows(content, table_width)
        if not rows:
            return {
                "status": "error",
                "message": "Could not parse content as table row(s). Use tab- or pipe-separated cell values, one row per line (or a single line for one row).",
                "block_id": block_id,
            }
        new_children: List[Dict[str, Any]] = []
        for row_cells in rows:
            cells = row_cells[:table_width]
            if len(cells) < table_width:
                cells.extend([""] * (table_width - len(cells)))
            new_children.append(
                {
                    "object": "block",
                    "type": "table_row",
                    "table_row": {
                        "cells": [
                            [
                                {
                                    "type": "text",
                                    "text": {"content": str(c).strip()},
                                    "annotations": _build_annotations(),
                                }
                            ]
                            for c in cells
                        ],
                    },
                }
            )
        _with_notion_rate_limit_retry(
            notion.blocks.children.append,
            block_id=block_id,
            children=new_children,
        )
        return {
            "status": "success",
            "message": f"Appended {len(new_children)} row(s) to table.",
            "block_id": block_id,
            "rows_appended": len(new_children),
        }

    def _resolve_parent_container(existing_block: Dict[str, Any]) -> str | None:
        """
        Resolve the container block/page id where replacement blocks should be appended.
        Uses the direct parent (page_id or block_id) to preserve nesting context.
        """
        parent = existing_block.get("parent", {}) or {}
        parent_type = parent.get("type")
        if parent_type == "page_id":
            return parent.get("page_id")
        if parent_type == "block_id":
            return parent.get("block_id")
        return None

    def _build_annotations() -> Dict[str, Any]:
        c = "default"
        if text_color is not None and str(text_color).strip():
            raw = str(text_color).strip().lower()
            if raw in NOTION_TEXT_COLORS:
                c = raw
        return {
            "bold": bool(bold) if bold is not None else False,
            "italic": bool(italic) if italic is not None else False,
            "underline": bool(underline) if underline is not None else False,
            "strikethrough": bool(strikethrough) if strikethrough is not None else False,
            "code": False,
            "color": c,
        }

    def _rich_text_for_update(text: str) -> List[Dict[str, Any]]:
        # LLMs often send bold:false, italic:false explicitly; treating "not None"
        # as uniform style skipped _rich_text_from_markup and left [text](url)
        # as literal characters in Notion.
        has_markup = _content_has_rich_markup(text)
        wants_uniform = (
            any(v is True for v in [bold, italic, underline, strikethrough])
            or (text_color is not None and str(text_color).strip() != "")
        )
        if has_markup:
            spans = _rich_text_from_markup(text)
            _merge_uniform_into_rich_text_spans(
                spans, bold, italic, underline, strikethrough, text_color
            )
            return spans
        if wants_uniform:
            return [
                {
                    "type": "text",
                    "text": {"content": text},
                    "annotations": _build_annotations(),
                }
            ]
        return _rich_text_from_markup(text)

    try:
        # Get the existing block to check its type and get parent info
        try:
            existing_block = _with_notion_rate_limit_retry(
                notion.blocks.retrieve, block_id=block_id
            )
        except Exception as retrieve_err:
            err_msg = str(retrieve_err).lower()
            # Notion can return 404 for table blocks when retrieved by id; append as table anyway
            if "404" in err_msg or "could not find" in err_msg or "not find" in err_msg:
                rows = _parse_table_content_to_rows(content, 0)
                if rows:
                    table_width = max(len(r) for r in rows) or 1
                    return _append_table_rows(block_id, content, table_width)
            raise retrieve_err

        block_type = existing_block.get("type")

        # Special handling for table blocks: append new row(s) to the table.
        # Content can be: tab-separated cells (one row), pipe-separated, or
        # markdown-style table lines (| a | b |). Parsed rows are appended as
        # table_row children.
        if block_type == "table":
            table_spec = existing_block.get("table") or {}
            table_width = int(table_spec.get("table_width", 0)) or 1
            # Parse content into rows of cell strings
            return _append_table_rows(block_id, content, table_width)

        # Special handling for table rows: allow updating the entire row by
        # providing a delimited string (tab- or pipe-separated) which is
        # converted into table_row.cells.
        if block_type == "table_row":
            # Existing cells to infer column count from
            table_row_content = existing_block.get("table_row", {}) or {}
            existing_cells = table_row_content.get("cells") or []
            col_count = len(existing_cells) if existing_cells else None

            # Split input into cell values. Prefer tabs, then pipes.
            if "\t" in content:
                raw_cells = [c.strip() for c in content.split("\t")]
            elif " | " in content:
                raw_cells = [c.strip() for c in content.split(" | ")]
            else:
                # Fallback: treat whole content as a single cell.
                raw_cells = [content.strip()]

            if col_count is None:
                col_count = len(raw_cells)

            # Pad or trim to match column count
            values = raw_cells[:col_count]
            if len(values) < col_count:
                values.extend([""] * (col_count - len(values)))

            new_cells: list[list[Dict[str, Any]]] = []
            for v in values:
                new_cells.append(
                    [
                        {
                            "type": "text",
                            "text": {"content": v},
                        }
                    ]
                )

            update_data = {"table_row": {"cells": new_cells}}
            _with_notion_rate_limit_retry(
                notion.blocks.update, block_id=block_id, **update_data
            )
            return {
                "status": "success",
                "message": "Table row updated successfully.",
                "block_id": block_id,
            }

        # Special handling for code blocks: keep native `code` block type
        # during edits instead of replacing it with paragraph/list blocks.
        if block_type == "code":
            code_content = existing_block.get("code", {}) or {}
            existing_language = code_content.get("language") or "plain text"
            if content is None:
                existing_rich_text = code_content.get("rich_text", []) or []
                content = "".join(
                    (item.get("plain_text") or item.get("text", {}).get("content", ""))
                    for item in existing_rich_text
                    if isinstance(item, dict)
                )
            update_data = {
                "code": {
                    "rich_text": [{"type": "text", "text": {"content": content or ""}}],
                    "language": existing_language,
                }
            }
            _with_notion_rate_limit_retry(
                notion.blocks.update, block_id=block_id, **update_data
            )
            return {
                "status": "success",
                "message": "Code block updated successfully.",
                "block_id": block_id,
                "language": existing_language,
            }

        block_content = existing_block.get(block_type, {}) if block_type else {}

        # If no content is provided, keep existing text and update formatting only.
        if content is None:
            existing_rich_text = block_content.get("rich_text", []) if isinstance(block_content, dict) else []
            content = "".join(
                (item.get("plain_text") or item.get("text", {}).get("content", ""))
                for item in existing_rich_text
                if isinstance(item, dict)
            )

        # For markdown-like multi-line content with list syntax, replace the
        # original block with native structured blocks so Notion renders UL/OL.
        if content is not None and all(
            v is None for v in [bold, italic, underline, strikethrough]
        ) and (text_color is None or not str(text_color).strip()):
            parsed_blocks = _blocks_from_text_content(content)
            existing_is_list = block_type in {"numbered_list_item", "bulleted_list_item"}
            # Preserve list semantics when updating an existing list item with
            # multiline content where first line has no explicit list marker.
            if (
                existing_is_list
                and parsed_blocks
                and parsed_blocks[0].get("type") == "paragraph"
            ):
                first_para = parsed_blocks[0].get("paragraph", {})
                converted_first = {
                    "object": "block",
                    "type": block_type,
                    block_type: {
                        "rich_text": first_para.get("rich_text", []),
                    },
                }
                # If the next parsed blocks are top-level bullets, attach them
                # as children of this converted list item so numbered-list
                # continuity is preserved (otherwise Notion restarts at 1).
                trailing_children: List[Dict[str, Any]] = []
                idx = 1
                while idx < len(parsed_blocks):
                    b = parsed_blocks[idx]
                    if b.get("type") == "bulleted_list_item":
                        trailing_children.append(b)
                        idx += 1
                        continue
                    break
                if trailing_children:
                    converted_first[block_type]["children"] = trailing_children
                    parsed_blocks = [converted_first] + parsed_blocks[idx:]
                else:
                    parsed_blocks[0] = converted_first

            # If first line has no explicit markup annotations, preserve the
            # existing block's first-span annotations for better style continuity.
            try:
                existing_rt = block_content.get("rich_text", []) if isinstance(block_content, dict) else []
                existing_ann = (
                    existing_rt[0].get("annotations")
                    if existing_rt and isinstance(existing_rt[0], dict)
                    else None
                )
                if existing_ann and parsed_blocks:
                    first_type = parsed_blocks[0].get("type")
                    first_payload = parsed_blocks[0].get(first_type, {}) if first_type else {}
                    first_rt = first_payload.get("rich_text", [])
                    for rt_item in first_rt:
                        if isinstance(rt_item, dict) and "annotations" not in rt_item:
                            rt_item["annotations"] = {
                                "bold": bool(existing_ann.get("bold")),
                                "italic": bool(existing_ann.get("italic")),
                                "underline": bool(existing_ann.get("underline")),
                                "strikethrough": bool(existing_ann.get("strikethrough")),
                                "code": bool(existing_ann.get("code", False)),
                                "color": existing_ann.get("color", "default"),
                            }
            except Exception:
                # Style preservation is best-effort; never block content updates.
                pass

            needs_structured_replace = (
                len(parsed_blocks) > 1
                or (len(parsed_blocks) == 1 and parsed_blocks[0].get("type") != "paragraph")
            )
            if needs_structured_replace:
                container_id = _resolve_parent_container(existing_block)
                if not container_id:
                    return {
                        "status": "error",
                        "message": "Unable to resolve parent container for structured block replacement.",
                        "block_id": block_id,
                    }
                _with_notion_rate_limit_retry(
                    notion.blocks.children.append,
                    block_id=container_id,
                    children=parsed_blocks,
                    after=block_id,
                )
                delete_warning = None
                try:
                    _with_notion_rate_limit_retry(notion.blocks.delete, block_id=block_id)
                except Exception as delete_err:
                    # Some block types/edge states can return 400 on delete even
                    # after replacement append succeeds. Try archive fallback and
                    # keep update successful to avoid false failure.
                    try:
                        _with_notion_rate_limit_retry(
                            notion.blocks.update, block_id=block_id, archived=True
                        )
                    except Exception as archive_err:
                        delete_warning = (
                            f"Replacement applied, but old block cleanup failed. "
                            f"delete_error={str(delete_err)} archive_error={str(archive_err)}"
                        )
                    else:
                        delete_warning = (
                            f"Replacement applied; delete returned error but archive fallback succeeded. "
                            f"delete_error={str(delete_err)}"
                        )
                result = {
                    "status": "success",
                    "message": "Block replaced with native Notion structured blocks (lists/paragraphs).",
                    "block_id": block_id,
                    "blocks_created": len(parsed_blocks),
                }
                if delete_warning:
                    result["warning"] = delete_warning
                return result

        # Only text-like blocks can be updated with a simple rich_text payload
        if "rich_text" not in block_content:
            if block_type == "child_page":
                return {
                    "status": "error",
                    "message": (
                        "The ID you provided is a page (database row), not a text block. "
                        "To change the page title use update_page_title. "
                        "To change date or other properties (e.g. a custom 'Created' or 'Date' field), "
                        "use update_page_properties with the page_id and the property names from get_database_schema. "
                        "Note: Notion's built-in 'Created time' cannot be changed via the API."
                    ),
                    "block_id": block_id,
                    "hint": "use_update_page_title_or_update_page_properties",
                }
            return {
                "status": "error",
                "message": f"Block type '{block_type}' does not support text updates via this tool.",
                "block_id": block_id,
            }

        # Notion has a 2000 character limit per text block
        MAX_CONTENT_LENGTH = 2000

        if len(content) <= MAX_CONTENT_LENGTH:
            # Simple case: content fits in one block
            update_data = {block_type: {"rich_text": _rich_text_for_update(content)}}
            _with_notion_rate_limit_retry(
                notion.blocks.update, block_id=block_id, **update_data
            )
            return {
                "status": "success",
                "message": "Block updated successfully.",
                "block_id": block_id,
            }
        else:
            # Content is too long - need to split into multiple blocks
            # Update the first block with first chunk
            first_chunk = content[:MAX_CONTENT_LENGTH]
            update_data = {block_type: {"rich_text": _rich_text_for_update(first_chunk)}}
            _with_notion_rate_limit_retry(
                notion.blocks.update, block_id=block_id, **update_data
            )

            # Get the parent to append remaining chunks
            parent = existing_block.get("parent", {})
            parent_type = parent.get("type")
            parent_id = None
            
            # Find the page_id - traverse up if parent is a block
            if parent_type == "page_id":
                parent_id = parent.get("page_id")
            elif parent_type == "block_id":
                # Traverse up to find the page
                current_parent_id = parent.get("block_id")
                max_depth = 10  # Safety limit
                depth = 0
                while depth < max_depth:
                    try:
                        parent_block = _with_notion_rate_limit_retry(
                            notion.blocks.retrieve, block_id=current_parent_id
                        )
                        parent_info = parent_block.get("parent", {})
                        if parent_info.get("type") == "page_id":
                            parent_id = parent_info.get("page_id")
                            break
                        elif parent_info.get("type") == "block_id":
                            current_parent_id = parent_info.get("block_id")
                            depth += 1
                        else:
                            break
                    except Exception:
                        break
            
            # Split remaining content into chunks
            remaining_content = content[MAX_CONTENT_LENGTH:]
            
            # Append remaining chunks as new blocks if we found a page_id
            if parent_id:
                new_blocks = _blocks_from_text_content(remaining_content)
                
                # Append new blocks to the parent page (they'll appear after the updated block)
                _with_notion_rate_limit_retry(
                    notion.blocks.children.append,
                    block_id=parent_id,
                    children=new_blocks,
                )
                
                return {
                    "status": "success", 
                    "message": f"Block updated successfully. Content was split into {len(new_blocks) + 1} blocks due to Notion's 2000 character limit per block.",
                    "block_id": block_id,
                    "blocks_created": len(new_blocks)
                }
            else:
                # If we can't determine the page, just update the first block and warn
                return {
                    "status": "partial_success",
                    "message": f"Block updated with first {MAX_CONTENT_LENGTH} characters. Remaining {len(remaining_content)} characters were truncated due to Notion's 2000 character limit per block. Please manually add the remaining content.",
                    "block_id": block_id,
                    "truncated": len(remaining_content)
                }
            
    except Exception as e:
        return {"status": "error", "message": f"Failed to update block: {str(e)}"}


def notion_delete_block(block_id: str = None, page_id: str = None, content_search: str = None, unified_token: str = None) -> Dict:
    """
    Delete a specific block from a Notion page.
    
    Args:
        block_id: Direct block ID to delete (preferred)
        page_id: Page ID to search within (required if block_id not provided)
        content_search: Text content to search for and delete (required if block_id not provided)
        unified_token: Authentication token
    """
    notion = get_notion_client(unified_token)
    
    try:
        # If block_id is provided, delete directly
        if block_id:
            _with_notion_rate_limit_retry(notion.blocks.delete, block_id=block_id)
            return {"status": "success", "message": "Block deleted successfully.", "block_id": block_id}
        
        # Otherwise, search for the block by content
        if not page_id or not content_search:
            return {"status": "error", "message": "Either block_id must be provided, or both page_id and content_search must be provided"}
        
        # Get all blocks from the page (paginate until we either find a match or exhaust)
        start_cursor = None
        while True:
            response = _with_notion_rate_limit_retry(
                notion.blocks.children.list, block_id=page_id, start_cursor=start_cursor
            )
            blocks = response.get("results", [])
            
            # Search for block containing the content
            for block in blocks:
                block_type = block.get("type")
                block_content = block.get(block_type, {})
                rich_text = block_content.get("rich_text", [])
                
                # Extract text content
                text_content = "".join(
                    [rt.get("text", {}).get("content", "") for rt in rich_text]
                )
                
                # Check if content matches (case-insensitive partial match)
                if content_search.lower() in text_content.lower():
                    _with_notion_rate_limit_retry(
                        notion.blocks.delete, block_id=block["id"]
                    )
                    return {
                        "status": "success", 
                        "message": f"Block containing '{content_search}' deleted successfully.",
                        "block_id": block["id"]
                    }

            if not response.get("has_more"):
                break
            start_cursor = response.get("next_cursor")
        
        return {"status": "error", "message": f"No block found containing '{content_search}'"}
        
    except Exception as e:
        return {"status": "error", "message": f"Failed to delete block: {str(e)}"}

def notion_get_comments(
    page_id: str,
    unified_token: str = None,
    comment_type: str = "page",
    block_id: str = None
) -> list:
    """
    Fetch comments for a Notion page, a specific block, or all blocks in a page.

    Args:
        page_id (str): The Notion page ID.
        unified_token (str, optional): The unified authentication token.
        type (str, optional): "page" (default) for page-level comments,
                              "block" for block-level comments.
        block_id (str, optional): If provided with type="block", fetches comments
                                  for that specific block only.

    Returns:
        list: List of comment objects (same structure as response.get("results", [])).
    """
    notion = get_notion_client(unified_token)
    target_type = str(comment_type).strip().lower() if comment_type else "page"
    print(f"Debug: Fetching '{target_type}' comments for page_id={page_id}")

    try:
        # Validate the page_id first
        notion.pages.retrieve(page_id)
    except Exception:
        return {"status": "error", "message": "Invalid page_id: Not a valid page or block."}
        return []

    try:
        if target_type == "page":
            # Case 1: Page-level comments
            response = _with_notion_rate_limit_retry(
                notion.comments.list, block_id=page_id
            )
            return response.get("results", [])

        elif target_type == "block":
            # Case 2: Block-level comments
            if block_id:
                # Specific block
                response = notion.comments.list(block_id=block_id)
                return response.get("results", [])
            else:
                # All blocks inside the page
                all_comments = []
                children: list[Dict[str, Any]] = []
                start_cursor = None
                while True:
                    child_resp = _with_notion_rate_limit_retry(
                        notion.blocks.children.list,
                        block_id=page_id,
                        start_cursor=start_cursor,
                    )
                    children.extend(child_resp.get("results", []))
                    if not child_resp.get("has_more"):
                        break
                    start_cursor = child_resp.get("next_cursor")

                for block in children:
                    try:
                        block_comments = _with_notion_rate_limit_retry(
                            notion.comments.list, block_id=block["id"]
                        ).get("results", [])
                        if block_comments:
                            all_comments.extend(block_comments)
                    except Exception as e:
                        print(f"Warning: Could not fetch comments for block {block['id']}: {e}")
                        continue
                return all_comments

        else:
            print(f"Warning: Unknown comment type '{target_type}'. Defaulting to page comments.")
            response = _with_notion_rate_limit_retry(
                notion.comments.list, block_id=page_id
            )
            return response.get("results", [])

    except Exception as e:
        print(f"Error fetching comments for {target_type}: {e}")
        return []


def add_notion_comment(
    page_id: str,
    content: str,
    unified_token: str = None,
    target_type: str = "page",
    block_id: str = None
) -> dict:
    """
    Add a comment to a Notion page or a specific block.

    Args:
        page_id (str): The Notion page ID.
        content (str): The text content of the comment.
        unified_token (str, optional): The unified authentication token.
        target_type (str, optional): "page" (default) or "block".
        block_id (str, optional): Specific block ID if commenting on a block.

    Returns:
        dict: A consistent response with status and message.
    """
    notion = get_notion_client(unified_token)
    resolved_type = str(target_type).strip().lower() if target_type else "page"
    print(f"Debug: Adding comment to {resolved_type} (page_id={page_id})")

    try:
        #  Validate the page first
        notion.pages.retrieve(page_id)
    except Exception:
        print(f"Error: Invalid page_id {page_id}")
        return {"status": "error", "message": "Invalid page_id: Not a valid page or block."}

    try:
        if resolved_type == "page":
            #  Add a comment directly to the page
            _with_notion_rate_limit_retry(
                notion.comments.create,
                parent={"block_id": page_id},
                rich_text=[{"type": "text", "text": {"content": content}}],
            )

        elif resolved_type == "block":
            # Add a comment to a specific block
            target_block_id = block_id if block_id else None
            if not target_block_id:
                return {"status": "error", "message": "Missing block_id for block comment."}

            _with_notion_rate_limit_retry(
                notion.comments.create,
                parent={"block_id": target_block_id},
                rich_text=[{"type": "text", "text": {"content": content}}],
            )

        else:
            print(f"Warning: Unknown target_type '{resolved_type}', defaulting to page comment.")
            _with_notion_rate_limit_retry(
                notion.comments.create,
                parent={"block_id": page_id},
                rich_text=[{"type": "text", "text": {"content": content}}],
            )

        return {"status": "success", "message": f"Comment added to {resolved_type}."}

    except Exception as e:
        print(f"Error adding comment to {resolved_type}: {e}")
        return {"status": "error", "message": str(e)}
    

def _plain_text_from_notion_block(block: Dict[str, Any], max_len: int = 2000) -> str:
    """Best-effort plain text for common Notion block shapes (API v2022+)."""
    if not isinstance(block, dict):
        return ""
    t = block.get("type")
    if not t:
        return ""
    payload = block.get(t, {}) or {}
    if t == "child_page":
        return (payload.get("title") or "").strip()[:max_len]
    if t in ("divider", "table_of_contents", "breadcrumb", "column_list", "column"):
        return f"[{t}]"
    if t == "table":
        return "[table]"
    if t == "table_row":
        cells = payload.get("cells") or []
        parts: List[str] = []
        for cell in cells:
            if not isinstance(cell, list):
                continue
            cell_text = "".join(
                (p.get("plain_text") or (p.get("text") or {}).get("content", ""))
                for p in cell
                if isinstance(p, dict)
            )
            parts.append(cell_text.strip())
        return " | ".join(parts)[:max_len]
    if t == "equation":
        expr = payload.get("expression")
        return (str(expr) if expr is not None else "")[:max_len]

    rt = payload.get("rich_text")
    if isinstance(rt, list) and rt:
        text = "".join(
            (p.get("plain_text") or (p.get("text") or {}).get("content", ""))
            for p in rt
            if isinstance(p, dict)
        )
        return text[:max_len]
    # Legacy to_do / paragraph sometimes used `text` instead of rich_text
    legacy = payload.get("text")
    if isinstance(legacy, list):
        text = "".join(
            (p.get("plain_text") or (p.get("text") or {}).get("content", ""))
            for p in legacy
            if isinstance(p, dict)
        )
        return text[:max_len]
    return ""


def notion_get_page_content(
    page_id: str,
    unified_token: str = None,
    max_blocks: int = 800,
) -> Dict[str, Any]:
    """
    Return a structured view of all blocks under a page (recursive).

    Each entry includes sibling ``index``, tree ``depth``, dotted ``path``
    (e.g. ``2/0/1``), ``block_id``, ``type``, and ``plain_text`` preview so
    tools and models can target the right block without parsing raw API JSON.
    """
    notion = get_notion_client(unified_token)
    flat: List[Dict[str, Any]] = []
    truncated = False

    def walk(parent_id: str, depth: int, path_prefix: str) -> None:
        nonlocal truncated
        if truncated or len(flat) >= max_blocks:
            truncated = True
            return
        start_cursor = None
        idx = 0
        while True:
            response = _with_notion_rate_limit_retry(
                notion.blocks.children.list,
                block_id=parent_id,
                start_cursor=start_cursor,
            )
            for block in response.get("results", []):
                if len(flat) >= max_blocks:
                    truncated = True
                    return
                bid = block.get("id")
                if not bid:
                    idx += 1
                    continue
                path = f"{path_prefix}/{idx}" if path_prefix else str(idx)
                entry: Dict[str, Any] = {
                    "block_id": bid,
                    "type": block.get("type"),
                    "index": idx,
                    "depth": depth,
                    "path": path,
                    "plain_text": _plain_text_from_notion_block(block),
                    "has_children": bool(block.get("has_children")),
                }
                if block.get("type") == "to_do":
                    td = block.get("to_do", {}) or {}
                    entry["checked"] = td.get("checked")
                if block.get("type") == "code":
                    cd = block.get("code", {}) or {}
                    lang = cd.get("language")
                    if lang:
                        entry["code_language"] = lang
                flat.append(entry)
                if block.get("has_children"):
                    walk(bid, depth + 1, path)
                idx += 1
            if not response.get("has_more"):
                break
            start_cursor = response.get("next_cursor")

    try:
        walk(page_id, 0, "")
    except Exception as e:
        return {
            "status": "error",
            "page_id": page_id,
            "message": str(e),
            "blocks": flat,
            "block_count": len(flat),
        }

    out: Dict[str, Any] = {
        "status": "success",
        "page_id": page_id,
        "block_count": len(flat),
        "blocks": flat,
    }
    if truncated:
        out["truncated"] = True
        out["message"] = (
            f"Block list truncated at {max_blocks} blocks. "
            "Narrow the page or use search if you need a specific section."
        )
    return out


def notion_update_page_title(page_id: str, new_title: str, unified_token: str = None):
    notion = get_notion_client(unified_token)
    page = _with_notion_rate_limit_retry(notion.pages.retrieve, page_id=page_id)
    props = page.get("properties", {})

    # Find the correct title property name dynamically
    title_prop_name = None
    for name, prop in props.items():
        if prop.get("type") == "title":
            title_prop_name = name
            break

    if not title_prop_name:
        return {"status": "error", "message": "No title property found on this page."}

    # Update the title
    _with_notion_rate_limit_retry(
        notion.pages.update,
        page_id=page_id,
        properties={title_prop_name: {"title": [{"text": {"content": new_title}}]}},
    )

    return {"status": "success", "message": f"Title updated to '{new_title}'"}


def _build_page_property_payload(
    prop_name: str,
    prop_info: Dict[str, Any],
    value: Any,
    notion,
) -> Optional[Dict[str, Any]]:
    """
    Build the Notion API payload for a single page property update.
    Returns None if the property type is not writable or value is invalid.
    """
    if value is None:
        return None
    prop_type = prop_info.get("type")
    if not prop_type:
        return None
    if prop_type == "title":
        return {"title": [{"text": {"content": str(value)}}]}
    if prop_type == "rich_text":
        return {"rich_text": [{"text": {"content": str(value)}}]}
    if prop_type == "number":
        try:
            return {"number": float(value) if value != "" else None}
        except (TypeError, ValueError):
            return None
    if prop_type == "checkbox":
        return {"checkbox": bool(value)}
    if prop_type == "select":
        options = prop_info.get("select", {}).get("options", [])
        option_names = {opt.get("name") for opt in options if isinstance(opt, dict) and opt.get("name")}
        if option_names and value not in option_names:
            return None
        return {"select": {"name": value}}
    if prop_type == "multi_select":
        if not isinstance(value, (list, tuple)):
            value = [value] if value is not None else []
        options = prop_info.get("multi_select", {}).get("options", [])
        option_names = {opt.get("name") for opt in options if isinstance(opt, dict) and opt.get("name")}
        if option_names:
            value = [v for v in value if v in option_names]
        return {"multi_select": [{"name": v} for v in value if v]}
    if prop_type == "status":
        options = prop_info.get("status", {}).get("options", [])
        option_names = {opt.get("name") for opt in options if isinstance(opt, dict) and opt.get("name")}
        if option_names and value not in option_names:
            return None
        return {"status": {"name": value}}
    if prop_type == "date":
        return {"date": {"start": value}}
    if prop_type == "people":
        people = []
        value_iter = [value] if isinstance(value, str) else (value or [])
        for person in value_iter:
            if isinstance(person, str):
                if "@" in person:
                    try:
                        uid = get_user_id_by_email(person, notion)
                        people.append({"id": uid})
                    except ValueError:
                        pass
                else:
                    people.append({"id": person})
        return {"people": people} if people else None
    if prop_type == "relation":
        rel_ids = value if isinstance(value, (list, tuple)) else [value]
        return {"relation": [{"id": r} for r in rel_ids if r]}
    if prop_type == "url":
        return {"url": value}
    if prop_type == "email":
        return {"email": value}
    if prop_type == "phone_number":
        return {"phone_number": value}
    return None


def notion_update_page_properties(
    page_id: str,
    properties: Dict[str, Any],
    unified_token: str = None,
) -> Dict[str, Any]:
    """
    Update one or more properties of a Notion page (e.g. a database row).
    Use get_database_schema to get exact property names and types.
    Note: Notion's built-in 'Created time' cannot be changed via the API.
    """
    if not properties:
        return {"status": "error", "message": "No properties provided."}
    notion = get_notion_client(unified_token)
    try:
        page = _with_notion_rate_limit_retry(notion.pages.retrieve, page_id=page_id)
    except Exception as e:
        return {"status": "error", "message": f"Failed to retrieve page: {str(e)}"}
    parent = page.get("parent", {}) or {}
    if parent.get("type") != "database_id":
        return {
            "status": "error",
            "message": "Page is not a database row; property updates are supported for database pages only.",
        }
    database_id = parent.get("database_id")
    if not database_id:
        return {
            "status": "error",
            "message": "Could not determine database for this page.",
        }

    schema = get_database_properties(database_id, notion)
    payload: Dict[str, Any] = {}
    read_only_props: list[str] = []

    for prop_name, value in properties.items():
        if prop_name not in schema:
            continue
        prop_info = schema[prop_name]
        prop_type = prop_info.get("type")

        # Explicitly reject known read-only / computed property types with
        # a clear error instead of silently dropping them.
        if prop_type in NON_WRITABLE_PROPERTY_TYPES:
            read_only_props.append(prop_name)
            continue

        one = _build_page_property_payload(prop_name, prop_info, value, notion)
        if one is not None:
            payload[prop_name] = one

    if not payload:
        if read_only_props:
            return {
                "status": "error",
                "message": (
                    "The following properties are read-only in Notion and "
                    "cannot be updated via the API: "
                    + ", ".join(sorted(set(read_only_props)))
                ),
            }
        return {
            "status": "error",
            "message": "No valid property updates; check property names (from get_database_schema) and value types.",
        }
    try:
        _with_notion_rate_limit_retry(
            notion.pages.update,
            page_id=page_id,
            properties=payload,
        )
        return {"status": "success", "message": "Page properties updated.", "updated": list(payload.keys())}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def notion_delete_page(page_id: str, unified_token: str = None) -> Dict:
    notion = get_notion_client(unified_token)
    _with_notion_rate_limit_retry(
        notion.pages.update, page_id=page_id, archived=True
    )
    return {"status": "success", "message": "Page archived"}


def notion_add_todo(page_id: str, todo_text: str, unified_token: str = None) -> Dict:
    notion = get_notion_client(unified_token)
    _with_notion_rate_limit_retry(
        notion.blocks.children.append,
        block_id=page_id,
        children=[
            {
                "object": "block",
                "type": "to_do",
                "to_do": {
                    "text": [{"type": "text", "text": {"content": todo_text}}],
                    "checked": False,
                },
            }
        ],
    )
    return {"status": "success", "message": "To-do added."}


def notion_query_database(
    database_id: str, status_value: str, unified_token: str = None
) -> List[Dict]:
    notion = get_notion_client(unified_token)
    # Inspect database schema to determine the correct status property dynamically
    properties = get_database_properties(database_id, notion)
    status_prop_name = find_status_property_name(properties)
    if not status_prop_name:
        raise ValueError(
            "No status property found in this database. Cannot filter by status."
        )

    pages: List[Dict[str, Any]] = []
    start_cursor = None

    while True:
        response = _with_notion_rate_limit_retry(
            notion.databases.query,
            database_id=database_id,
            filter={"property": status_prop_name, "status": {"equals": status_value}},
            start_cursor=start_cursor,
        )
        results = response.get("results", [])
        for r in results:
            pages.append({"id": r.get("id"), "url": r.get("url")})

        if not response.get("has_more"):
            break
        start_cursor = response.get("next_cursor")

    return pages


def notion_list_all_pages(unified_token: str = None):
    notion = get_notion_client(unified_token)
    pages: List[Dict[str, Any]] = []
    start_cursor = None

    while True:
        response = _with_notion_rate_limit_retry(
            notion.search,
            filter={"property": "object", "value": "page"},
            start_cursor=start_cursor,
        )
        results = response.get("results", [])

        for p in results:
            props = p.get("properties", {}) or {}
            title = get_title_from_properties(props)
            pages.append({"id": p.get("id"), "title": title, "url": p.get("url")})

        if not response.get("has_more"):
            break
        start_cursor = response.get("next_cursor")

    return pages


def notion_get_documents(unified_token: str = None):
    notion = get_notion_client(unified_token)
    documents: List[Dict[str, Any]] = []
    start_cursor = None

    while True:
        response = _with_notion_rate_limit_retry(
            notion.search,
            filter={"property": "object", "value": "page"},
            start_cursor=start_cursor,
        )
        results = response.get("results", [])

        for page in results:
            props = page.get("properties", {}) or {}
            title = get_title_from_properties(props)

            last_edited_time = page.get("last_edited_time")
            editor_info = page.get("last_edited_by", {}) or {}
            editor_name = "Unknown"

            if editor_info.get("object") == "user":
                editor_name = editor_info.get("name") or editor_info.get("id", "Unknown")

            item = {
                "id": page.get("id"),
                "title": title,
                "url": page.get("url"),
                "last_edited_time": last_edited_time,
                "type": "Notion",
                "last_edited_by": editor_name,
            }

            # Treat pages that do not look like task rows as generic documents
            if "Task name" not in props:
                documents.append(item)

        if not response.get("has_more"):
            break
        start_cursor = response.get("next_cursor")

    return {"documents": documents}


def get_database_title(unified_token: str = None, database_id: str = None):
    try:
        notion = get_notion_client(unified_token)
        db = _with_notion_rate_limit_retry(
            notion.databases.retrieve, database_id=database_id
        )
        title = db.get("title", [])
        if title:
            return "".join([t.get("text", {}).get("content", "") for t in title])
        return None
    except Exception as e:
        print(f"Error retrieving database title: {e}")
        return None


def notion_get_tasks_detailed(unified_token: str = None):
    notion = get_notion_client(unified_token)
    tasks: List[Dict[str, Any]] = []

    # Cache database schemas so we don't repeatedly call databases.retrieve
    db_config_cache: Dict[str, Dict[str, Any]] = {}

    def _infer_task_property_config(db_id: str) -> Dict[str, Any]:
        """Infer which properties in a database represent common task fields."""
        if db_id in db_config_cache:
            return db_config_cache[db_id]

        db = _with_notion_rate_limit_retry(
            notion.databases.retrieve, database_id=db_id
        )
        db_props = db.get("properties", {}) or {}

        # Database title (for related_to field)
        db_title_parts = db.get("title", []) or []
        db_title_pieces: list[str] = []
        for part in db_title_parts:
            if "plain_text" in part:
                db_title_pieces.append(part.get("plain_text") or "")
            elif "text" in part:
                db_title_pieces.append(part["text"].get("content", ""))
        database_title = "".join(db_title_pieces).strip() or None

        # Title
        title_prop_name = get_title_property_name(db_props)

        # Status
        status_prop_name = find_status_property_name(db_props)

        # Priority: select/multi_select with 'priority' in the name, or first such property
        priority_prop_name = None
        for name, prop in db_props.items():
            if prop.get("type") in ("select", "multi_select") and "priority" in name.lower():
                priority_prop_name = name
                break
        if not priority_prop_name:
            for name, prop in db_props.items():
                if prop.get("type") in ("select", "multi_select"):
                    priority_prop_name = name
                    break

        # Due date: date property with common labels
        due_prop_name = None
        for name, prop in db_props.items():
            if prop.get("type") == "date":
                lower = name.lower()
                if any(k in lower for k in ["due", "deadline", "target", "date"]):
                    due_prop_name = name
                    break

        # Stage: rich_text or select property containing 'stage'
        stage_prop_name = None
        for name, prop in db_props.items():
            if prop.get("type") in ("rich_text", "select", "multi_select") and "stage" in name.lower():
                stage_prop_name = name
                break

        # Type: property containing 'type'
        type_prop_name = None
        for name, prop in db_props.items():
            if "type" in name.lower() and prop.get("type") in (
                "select",
                "multi_select",
                "rich_text",
            ):
                type_prop_name = name
                break

        # Category: property containing 'category'
        category_prop_name = None
        for name, prop in db_props.items():
            if "category" in name.lower() and prop.get("type") in (
                "select",
                "multi_select",
                "rich_text",
            ):
                category_prop_name = name
                break

        # Reference: property containing 'reference' or 'ref'
        reference_prop_name = None
        for name, prop in db_props.items():
            lower = name.lower()
            if any(k in lower for k in ["reference", "ref"]) and prop.get("type") in (
                "url",
                "rich_text",
                "relation",
            ):
                reference_prop_name = name
                break

        # Assignee: people property
        assignee_prop_name = find_people_property_name(db_props)

        # Description: first rich_text property that isn't clearly a small label
        description_prop_name = None
        for name, prop in db_props.items():
            if prop.get("type") == "rich_text" and name not in (
                stage_prop_name,
                reference_prop_name,
            ):
                description_prop_name = name
                break

        config = {
            "title": title_prop_name,
            "status": status_prop_name,
            "priority": priority_prop_name,
            "due_date": due_prop_name,
            "stage": stage_prop_name,
            "type": type_prop_name,
            "category": category_prop_name,
            "reference": reference_prop_name,
            "assignee": assignee_prop_name,
            "description": description_prop_name,
            "database_title": database_title,
        }
        db_config_cache[db_id] = config
        return config

    # Search all pages with pagination and treat pages that belong to databases
    # (and have a status field) as tasks.
    start_cursor = None
    while True:
        response = _with_notion_rate_limit_retry(
            notion.search,
            filter={"property": "object", "value": "page"},
            start_cursor=start_cursor,
        )
        results = response.get("results", [])

        for page in results:
            parent = page.get("parent", {}) or {}
            if parent.get("type") != "database_id":
                continue

            database_id = parent.get("database_id")
            if not database_id:
                continue

            config = _infer_task_property_config(database_id)
            # Skip databases that don't look like task trackers (no status field)
            if not config.get("status"):
                continue

            props = page.get("properties", {}) or {}

            # Title
            if config.get("title") and config["title"] in props:
                title_val = extract_property_value(props.get(config["title"], {}))
                title = title_val or "(no title)"
            else:
                title = get_title_from_properties(props)

            task_data: Dict[str, Any] = {
                "id": page.get("id"),
                "title": title,
                "url": page.get("url"),
                "stage": extract_property_value(
                    props.get(config.get("stage")) if config.get("stage") else {}
                ),
                "type": extract_property_value(
                    props.get(config.get("type")) if config.get("type") else {}
                ),
                "status": extract_property_value(
                    props.get(config.get("status")) if config.get("status") else {}
                ),
                "priority": extract_property_value(
                    props.get(config.get("priority")) if config.get("priority") else {}
                ),
                "due_date": extract_property_value(
                    props.get(config.get("due_date")) if config.get("due_date") else {}
                ),
                "source": "Notion",
                "reference": extract_property_value(
                    props.get(config.get("reference"))
                    if config.get("reference")
                    else {}
                ),
                # Keep raw people objects for backward compatibility
                "assignee": (
                    props.get(config.get("assignee"), {}).get("people", [])
                    if config.get("assignee")
                    else []
                ),
                "related_to": config.get("database_title"),
                "database_id": database_id,
                "description": extract_property_value(
                    props.get(config.get("description"))
                    if config.get("description")
                    else {}
                ),
                "category": extract_property_value(
                    props.get(config.get("category"))
                    if config.get("category")
                    else {}
                ),
                "last_modified": page.get("last_edited_time"),
                "author": "Unknown",
            }

            editor_info = page.get("last_edited_by", {}) or {}
            if editor_info.get("object") == "user":
                # Use embedded name when available (no extra API calls)
                task_data["author"] = editor_info.get("name") or editor_info.get(
                    "id", "Unknown"
                )

            tasks.append(task_data)

        if not response.get("has_more"):
            break
        start_cursor = response.get("next_cursor")

    return {"tasks": tasks}


def notion_find_tasks_tracker_db(unified_token: str = None) -> dict:
    """Find a Notion database titled 'Tasks Tracker' (case-insensitive)."""
    notion = get_notion_client(unified_token)
    results = notion.search(filter={"property": "object", "value": "database"}).get(
        "results", []
    )
    for db in results:
        title_parts = db.get("title", [])
        title = "".join([t.get("text", {}).get("content", "") for t in title_parts])
        if title and title.strip().lower() == "tasks tracker":
            return {"id": db.get("id"), "title": title}
    return {}


def get_database_properties(database_id, notion):
    db = _with_notion_rate_limit_retry(
        notion.databases.retrieve, database_id=database_id
    )
    return db.get("properties", {})


def get_title_property_name(properties):
    for prop_name, prop_info in properties.items():
        if prop_info.get("type") == "title":
            return prop_name
    return None  # fallback if none found


def find_status_property_name(properties: Dict[str, Any]) -> str | None:
    """
    Find the most appropriate status property in a database schema.
    Prefer properties whose type is 'status', otherwise fall back to a
    property literally named 'Status' (case-insensitive).
    """
    for prop_name, prop_info in properties.items():
        if prop_info.get("type") == "status":
            return prop_name
    for prop_name in properties.keys():
        if prop_name.lower() == "status":
            return prop_name
    return None


def find_people_property_name(properties: Dict[str, Any]) -> str | None:
    """
    Find the most appropriate people property in a database schema.
    Prefer properties whose type is 'people' and whose name suggests assignee/owner.
    """
    # First pass: semantic match on common assignee/owner labels
    preferred_keywords = ["assignee", "owner", "responsible", "assigned"]
    for prop_name, prop_info in properties.items():
        if prop_info.get("type") == "people":
            lower = prop_name.lower()
            if any(keyword in lower for keyword in preferred_keywords):
                return prop_name

    # Fallback: first people-type property
    for prop_name, prop_info in properties.items():
        if prop_info.get("type") == "people":
            return prop_name
    return None


def notion_create_task(
    unified_token: str = None, database_id: str = None, data: dict = None
):
    notion = get_notion_client(unified_token)
    data = data or {}
    # If database_id not provided, default to user's 'tasks_tracker' DB if available
    if not database_id:
        tt = notion_find_tasks_tracker_db(unified_token)
        database_id = tt.get("id") if tt else None
    if not database_id:
        return {
            "success": False,
            "error": "No database_id provided and tasks_tracker not found",
        }

    properties = get_database_properties(database_id, notion)
    title_property_name = get_title_property_name(properties)
    if not title_property_name:
        return {"success": False, "error": "No title property found in the database"}

    # Normalize alias keys
    input_aliases = {"assignee_ids": "assignee", "related_to_ids": "related_to"}
    for alias, canonical in input_aliases.items():
        if alias in data and canonical not in data:
            data[canonical] = data[alias]

    def get_property_value(prop_name, value):
        if prop_name not in properties:
            return None

        prop_type = properties[prop_name]["type"]
        if value is None:
            return None

        if prop_type == "title":
            return {"title": [{"text": {"content": value}}]}
        elif prop_type == "rich_text":
            return {"rich_text": [{"text": {"content": value}}]}
        elif prop_type == "select":
            # Validate against known options when available
            options = properties[prop_name].get("select", {}).get("options", [])
            option_names = {
                opt.get("name")
                for opt in options
                if isinstance(opt, dict) and opt.get("name") is not None
            }
            if option_names and value not in option_names:
                raise ValueError(
                    f"Invalid value '{value}' for select property '{prop_name}'. "
                    f"Valid options are: {sorted(option_names)}"
                )
            return {"select": {"name": value}}
        elif prop_type == "multi_select":
            if not isinstance(value, (list, tuple)):
                raise ValueError(
                    f"multi_select property '{prop_name}' expects a list of option names."
                )
            options = properties[prop_name].get("multi_select", {}).get(
                "options", []
            )
            option_names = {
                opt.get("name")
                for opt in options
                if isinstance(opt, dict) and opt.get("name") is not None
            }
            invalid = [v for v in value if option_names and v not in option_names]
            if invalid:
                raise ValueError(
                    f"Invalid values {invalid} for multi_select property '{prop_name}'. "
                    f"Valid options are: {sorted(option_names)}"
                )
            return {"multi_select": [{"name": v} for v in value if v]}
        elif prop_type == "status":
            options = properties[prop_name].get("status", {}).get("options", [])
            option_names = {
                opt.get("name")
                for opt in options
                if isinstance(opt, dict) and opt.get("name") is not None
            }
            if option_names and value not in option_names:
                raise ValueError(
                    f"Invalid status '{value}' for property '{prop_name}'. "
                    f"Valid options are: {sorted(option_names)}"
                )
            return {"status": {"name": value}}
        elif prop_type == "date":
            return {"date": {"start": value}}
        # elif prop_type == "people":
        #     return {
        #         "people": [{"id": v} for v in value if v]
        #     }

        elif prop_type == "people":
            people = []
            # Accept a single string or a list of strings
            if isinstance(value, str):
                value_iter = [value]
            else:
                value_iter = value
            for person in value_iter:
                if isinstance(person, str):
                    if "@" in person:
                        try:
                            uid = get_user_id_by_email(person, notion)
                            people.append({"id": uid})
                        except ValueError as e:
                            print(f"[Warning] {e}")
                    else:
                        # Assume it's a user ID
                        people.append({"id": person})
            return {"people": people}

        elif prop_type == "relation":
            return {"relation": [{"id": v} for v in value if v]}
        elif prop_type == "url":
            return {"url": value}
        return None

    # Heuristics to infer which database properties correspond to common task fields
    status_prop_name = find_status_property_name(properties)

    # Priority: select/multi_select with 'priority' in the name, or first such property
    priority_prop_name = None
    for name, prop in properties.items():
        if prop.get("type") in ("select", "multi_select") and "priority" in name.lower():
            priority_prop_name = name
            break
    if not priority_prop_name:
        for name, prop in properties.items():
            if prop.get("type") in ("select", "multi_select"):
                priority_prop_name = name
                break

    # Due date: date property with common labels
    due_date_prop_name = None
    for name, prop in properties.items():
        if prop.get("type") == "date":
            lower = name.lower()
            if any(k in lower for k in ["due", "deadline", "target", "date"]):
                due_date_prop_name = name
                break

    # Stage: rich_text or select property containing 'stage'
    stage_prop_name = None
    for name, prop in properties.items():
        if prop.get("type") in ("rich_text", "select", "multi_select") and "stage" in name.lower():
            stage_prop_name = name
            break

    # Type: property containing 'type'
    type_prop_name = None
    for name, prop in properties.items():
        if "type" in name.lower() and prop.get("type") in (
            "select",
            "multi_select",
            "rich_text",
        ):
            type_prop_name = name
            break

    # Category: property containing 'category'
    category_prop_name = None
    for name, prop in properties.items():
        if "category" in name.lower() and prop.get("type") in (
            "select",
            "multi_select",
            "rich_text",
        ):
            category_prop_name = name
            break

    # Reference: property containing 'reference' or 'ref'
    reference_prop_name = None
    for name, prop in properties.items():
        lower = name.lower()
        if any(k in lower for k in ["reference", "ref"]) and prop.get("type") in (
            "url",
            "rich_text",
            "relation",
        ):
            reference_prop_name = name
            break

    # Assignee: people property
    assignee_prop_name = find_people_property_name(properties)

    # Related-to: relation property
    related_to_prop_name = None
    for name, prop in properties.items():
        if prop.get("type") == "relation":
            lower = name.lower()
            if any(k in lower for k in ["related", "relation", "linked", "links"]):
                related_to_prop_name = name
                break
    if not related_to_prop_name:
        for name, prop in properties.items():
            if prop.get("type") == "relation":
                related_to_prop_name = name
                break

    # Description: first rich_text property that isn't clearly a small label
    description_prop_name = None
    for name, prop in properties.items():
        if prop.get("type") == "rich_text" and name not in (
            stage_prop_name,
            reference_prop_name,
        ):
            description_prop_name = name
            break

    semantic_to_prop = {
        "title": title_property_name,
        "description": description_prop_name,
        "status": status_prop_name,
        "priority": priority_prop_name,
        "due_date": due_date_prop_name,
        "assignee": assignee_prop_name,
        "reference": reference_prop_name,
        "type": type_prop_name,
        "related_to": related_to_prop_name,
        "stage": stage_prop_name,
        "category": category_prop_name,
    }

    payload = {"parent": {"database_id": database_id}, "properties": {}}

    try:
        # 1) Direct property-name keys (allow caller to specify exact Notion property names)
        for prop_name in properties.keys():
            if prop_name in data:
                prop_value = get_property_value(prop_name, data[prop_name])
                if prop_value is not None:
                    payload["properties"][prop_name] = prop_value

        # 2) Semantic keys mapped to inferred properties (only if not already set above)
        for semantic_key, prop_name in semantic_to_prop.items():
            if not prop_name or semantic_key not in data:
                continue
            if prop_name in payload["properties"]:
                continue
            prop_value = get_property_value(prop_name, data[semantic_key])
            if prop_value is not None:
                payload["properties"][prop_name] = prop_value

        created = _with_notion_rate_limit_retry(notion.pages.create, **payload)
        return {"success": True, "id": created.get("id"), "url": created.get("url")}
    except Exception as e:
        return {"success": False, "error": str(e)}


def get_title_field_from_page(page_id, notion):
    page = _with_notion_rate_limit_retry(notion.pages.retrieve, page_id=page_id)
    for key, value in page["properties"].items():
        if value["type"] == "title":
            return key
    raise ValueError("No title field found in the page.")


def get_database_id_from_page(page_id, notion):
    page = _with_notion_rate_limit_retry(notion.pages.retrieve, page_id=page_id)
    return page["parent"]["database_id"]


def get_status_field_and_options_from_database(page_id, notion):
    database_id = get_database_id_from_page(page_id, notion)
    db = _with_notion_rate_limit_retry(
        notion.databases.retrieve, database_id=database_id
    )
    properties = db.get("properties", {}) or {}
    status_field = find_status_property_name(properties)
    if not status_field:
        return None, []

    status_prop = properties.get(status_field, {})
    options = [
        opt.get("name")
        for opt in status_prop.get("status", {}).get("options", [])
        if isinstance(opt, dict) and opt.get("name") is not None
    ]
    return status_field, options


def get_rich_text_field_from_page(page_id, notion):
    page = _with_notion_rate_limit_retry(notion.pages.retrieve, page_id=page_id)
    for key, value in page["properties"].items():
        if value["type"] == "rich_text":
            if key == "Stage":
                continue
            return key
    raise ValueError("No rich_text field found in the page.")


def edit_task(unified_token: str = None, task_id: str = None, updates: dict = None):
    notion = get_notion_client(unified_token)
    props = {}

    title_field = (
        get_title_field_from_page(task_id, notion) if "title" in updates else None
    )
    status_field, status_options = (
        get_status_field_and_options_from_database(task_id, notion)
        if "status" in updates
        else (None, [])
    )
    # description_field = get_rich_text_field_from_page(task_id) if "description" in updates else None

    for key, value in updates.items():
        if key == "title":
            props[title_field] = {"title": [{"text": {"content": value}}]}
        elif key == "due_date":
            props["Due date"] = {"date": {"start": value}}
        elif key == "priority":
            props["Priority"] = {"select": {"name": value}}
        elif key == "status":
            if value not in status_options:
                raise ValueError(
                    f"Invalid status: '{value}'. Valid options are: {status_options}"
                )
            props[status_field] = {"status": {"name": value}}
        elif key == "description":
            props["Description"] = {"rich_text": [{"text": {"content": value}}]}
        elif key == "stage":
            props["Stage"] = {"rich_text": [{"text": {"content": value}}]}
        # elif key == "assignee":
        #     assignees = value if isinstance(value, list) else [value]
        #     props["Assignee"] = {
        #         "people": [{"object": "user", "id": uid} for uid in assignees]
        #     }
        elif key == "assignee":
            emails = value if isinstance(value, list) else [value]
            props["Assignee"] = {
                "people": [
                    {"object": "user", "id": get_user_id_by_email(email, notion)}
                    for email in emails
                ]
            }

    return notion.pages.update(page_id=task_id, properties=props)


def get_user_id_by_email(email, notion):
    """
    Resolve a Notion user ID from an email address, with basic in-memory caching
    and pagination over the users.list API.
    """
    if not email:
        raise ValueError("Email must be provided to look up a Notion user.")

    cached_id = _USER_EMAIL_CACHE.get(email.lower())
    if cached_id:
        return cached_id

    start_cursor = None
    while True:
        response = _with_notion_rate_limit_retry(
            notion.users.list, start_cursor=start_cursor
        )
        for user in response.get("results", []):
            if user.get("type") != "person":
                continue
            person_info = user.get("person", {}) or {}
            user_email = (person_info.get("email") or "").lower()
            if not user_email:
                continue
            user_id = user.get("id")
            if user_id:
                _USER_EMAIL_CACHE[user_email] = user_id
            if user_email == email.lower():
                return user_id

        if not response.get("has_more"):
            break
        start_cursor = response.get("next_cursor")

    raise ValueError(f"No Notion user found with email: {email}")


def mark_task_complete(unified_token: str = None, task_id: str = None):
    notion = get_notion_client(unified_token)
    status_field, status_options = get_status_field_and_options_from_database(
        task_id, notion
    )
    if not status_field:
        raise ValueError(
            "No status property found for this task's database; cannot mark complete."
        )

    # Prefer an existing option named 'Completed' (case-insensitive) if available
    target_status = next(
        (opt for opt in status_options if opt and opt.lower() == "completed"),
        "Completed",
    )

    return _with_notion_rate_limit_retry(
        notion.pages.update,
        page_id=task_id,
        properties={status_field: {"status": {"name": target_status}}},
    )


def delegate_task(
    unified_token: str = None, task_id: str = None, new_assignee_id: str = None
):
    notion = get_notion_client(unified_token)
    if not new_assignee_id:
        raise ValueError("new_assignee_id must be provided to delegate a task.")

    # Resolve the database and choose an appropriate people property dynamically
    database_id = get_database_id_from_page(task_id, notion)
    properties = get_database_properties(database_id, notion)
    people_field = find_people_property_name(properties)
    if not people_field:
        raise ValueError(
            "No people-type property found in this task's database; cannot delegate."
        )

    return _with_notion_rate_limit_retry(
        notion.pages.update,
        page_id=task_id,
        properties={
            people_field: {"people": [{"object": "user", "id": new_assignee_id}]}
        },
    )


def delete_task(unified_token: str = None, task_id: str = None):
    notion = get_notion_client(unified_token)
    return notion.pages.update(page_id=task_id, archived=True)


def notion_get_users(unified_token: str = None):
    notion = get_notion_client(unified_token)

    try:
        users = []
        start_cursor = None

        while True:
            response = _with_notion_rate_limit_retry(
                notion.users.list, start_cursor=start_cursor
            )
            for user in response.get("results", []):
                if user["type"] == "person":
                    users.append(
                        {
                            "label": user.get("name"),
                            "value": user.get("id"),
                            "email": user.get("person", {}).get("email"),
                            "avatar_url": user.get("avatar_url"),
                        }
                    )

            if response.get("has_more"):
                start_cursor = response.get("next_cursor")
            else:
                break

        return {"success": True, "users": users}

    except Exception as e:
        return {"success": False, "error": str(e)}


def get_all_blocks(notion, block_id):
    """Recursively fetch all child blocks under a given block"""
    blocks = []
    queue = [block_id]

    while queue:
        current = queue.pop()
        try:
            start_cursor = None
            while True:
                response = _with_notion_rate_limit_retry(
                    notion.blocks.children.list,
                    block_id=current,
                    start_cursor=start_cursor,
                )
                children = response.get("results", [])
                blocks.extend(children)
                for child in children:
                    if child.get("has_children"):
                        queue.append(child["id"])
                if not response.get("has_more"):
                    break
                start_cursor = response.get("next_cursor")
        except Exception as e:
            print(f"Error fetching children for block {current}: {e}")
    return blocks


def get_comments_on_page(unified_token: str = None, page_id: str = None):
    notion = get_notion_client(unified_token)

    try:
        # Start with the top-level page block
        all_blocks = [{"id": page_id}] + get_all_blocks(notion, page_id)
        comments_data = []

        for block in all_blocks:
            block_id = block["id"]
            try:
                comments = _with_notion_rate_limit_retry(
                    notion.comments.list, block_id=block_id
                ).get("results", [])
                if comments:
                    comments_data.append({"block_id": block_id, "comments": comments})
            except Exception as e:
                print(f"Error fetching comments for block {block_id}: {e}")

        return comments_data

    except Exception as e:
        print("Error:", e)
        return []


def is_notion_key_valid(notion_key: str) -> bool:
    # Here notion_key is expected to be a direct Notion API key, not a unified token
    notion = Client(auth=notion_key)
    try:
        notion.users.me()  # this fails if the token is invalid
        return True
    except Exception as e:
        print(f"Invalid Notion token: {e}")
        return False


def blocks_to_markdown(blocks):
    md = []
    for block in blocks:
        t = block.get("type")
        b = block.get(t, {})
        text = "".join([r.get("plain_text", "") for r in b.get("rich_text", [])])

        if t == "heading_1":
            md.append("# " + text)
        elif t == "heading_2":
            md.append("## " + text)
        elif t == "heading_3":
            md.append("### " + text)
        elif t == "bulleted_list_item":
            md.append("- " + text)
        elif t == "numbered_list_item":
            md.append("1. " + text)
        elif t == "to_do":
            checked = "[x]" if b.get("checked") else "[ ]"
            md.append(f"- {checked} {text}")
        else:
            md.append(text)

    return "\n".join(md)


def blocks_to_docx(blocks):
    doc = Document()

    for block in blocks:
        t = block.get("type")
        b = block.get(t, {})
        text = "".join([r.get("plain_text", "") for r in b.get("rich_text", [])])

        if not text:
            continue

        if t == "heading_1":
            doc.add_heading(text, level=1)
        elif t == "heading_2":
            doc.add_heading(text, level=2)
        elif t == "heading_3":
            doc.add_heading(text, level=3)
        elif t in ["bulleted_list_item", "numbered_list_item"]:
            doc.add_paragraph(text, style="List Bullet")
        elif t == "to_do":
            checked = "☑" if b.get("checked") else "☐"
            doc.add_paragraph(f"{checked} {text}")
        else:
            doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def download_notion_export(
    unified_token: str = None, page_id: str = None, fmt: str = None
):
    notion = get_notion_client(unified_token)

    try:
        # Fetch all blocks
        blocks = []
        start_cursor = None
        while True:
            response = _with_notion_rate_limit_retry(
                notion.blocks.children.list,
                block_id=page_id,
                start_cursor=start_cursor,
            )
            blocks.extend(response.get("results", []))
            if not response.get("has_more"):
                break
            start_cursor = response.get("next_cursor")

        # Convert to desired format
        if fmt == "md":
            md_content = blocks_to_markdown(blocks)
            return md_content, "text/markdown", f"{page_id}.md"
        elif fmt == "docx":
            docx_data = blocks_to_docx(blocks)
            return (
                docx_data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{page_id}.docx",
            )

        else:
            return None, "text/plain", "unsupported.txt"

    except Exception as e:
        print("Error in download_notion_export:", e)
        return None, "text/plain", "error.txt"


# === HTML to Notion Block Conversion Functions ===

MAX_BLOCKS_PER_REQUEST = 100


def clean_id(id_str: str) -> str:
    """Ensure valid UUID format for Notion block IDs"""
    return str(uuid.UUID(re.sub(r"[^a-fA-F0-9]", "", id_str)))


def html_to_notion_blocks(html: str) -> List[Dict]:
    soup = BeautifulSoup(html, "html.parser")
    blocks = []

    for elem in soup.body or soup:
        if isinstance(elem, str):
            continue
        tag = elem.name
        text = elem.get_text(strip=True)

        # Handle paragraphs and headers
        if tag == "p":
            blocks.append(paragraph_block(text))
        elif tag in ["h1", "h2", "h3"]:
            blocks.append(heading_block(text, level=int(tag[1])))
        elif tag == "ul":
            blocks.extend(bulleted_list_block(elem))
        elif tag == "ol":
            blocks.extend(numbered_list_block(elem))
        elif tag == "img":
            blocks.append(image_block(elem))
        elif tag == "table":
            blocks.append(table_block(elem))
        elif tag == "pre" or tag == "code":
            blocks.append(code_block(elem))
        else:
            blocks.append(paragraph_block(f"[UNSUPPORTED HTML: <{tag}>] {text}"))

    return blocks


def paragraph_block(text: str) -> Dict:
    return {
        "object": "block",
        "type": "paragraph",
        "paragraph": {
            "rich_text": [{"type": "text", "text": {"content": text[:2000]}}]
        }
    }


def heading_block(text: str, level: int) -> Dict:
    heading_type = f"heading_{min(level,3)}"
    return {
        "object": "block",
        "type": heading_type,
        heading_type: {
            "rich_text": [{"type": "text", "text": {"content": text[:2000]}}]
        }
    }


def bulleted_list_block(ul_tag) -> List[Dict]:
    return [
        {
            "object": "block",
            "type": "bulleted_list_item",
            "bulleted_list_item": {
                "rich_text": [{"type": "text", "text": {"content": li.get_text(strip=True)[:2000]}}]
            }
        }
        for li in ul_tag.find_all("li")
    ]


def numbered_list_block(ol_tag) -> List[Dict]:
    return [
        {
            "object": "block",
            "type": "numbered_list_item",
            "numbered_list_item": {
                "rich_text": [{"type": "text", "text": {"content": li.get_text(strip=True)[:2000]}}]
            }
        }
        for li in ol_tag.find_all("li")
    ]


def image_block(img_tag) -> Dict:
    src = img_tag.get("src", "")
    return {
        "object": "block",
        "type": "image",
        "image": {
            "type": "external",
            "external": {"url": src}
        }
    }


def code_block(tag) -> Dict:
    code_text = tag.get_text(strip=True)[:2000]
    return {
        "object": "block",
        "type": "code",
        "code": {
            "rich_text": [{"type": "text", "text": {"content": code_text}}],
            "language": "plain text"
        }
    }


def table_block(table_tag) -> Dict:
    rows = []
    for row_tag in table_tag.find_all("tr"):
        cells = row_tag.find_all(["td", "th"])
        cell_texts = [cell.get_text(strip=True)[:2000] for cell in cells]
        rows.append(" | ".join(cell_texts))

    table_string = "\n".join(rows)
    return code_block(BeautifulSoup(table_string, "html.parser"))  # fallback as code block


# === Page Creation with HTML Support ===

def create_notion_page(unified_token: str, title: str, html_string: str, database_id: str) -> Dict:
    notion = get_notion_client(unified_token)

    all_blocks = html_to_notion_blocks(html_string)
    block_chunks = [all_blocks[i:i + MAX_BLOCKS_PER_REQUEST] for i in range(0, len(all_blocks), MAX_BLOCKS_PER_REQUEST)]

    try:
        page = _with_notion_rate_limit_retry(
            notion.pages.create,
            parent={"type": "database_id", "database_id": database_id},
            properties={
                "title": {
                    "title": [
                        {"type": "text", "text": {"content": title[:2000]}}
                    ]
                }
            },
            children=block_chunks[0] if block_chunks else [],
        )

        for chunk in block_chunks[1:]:
            _with_notion_rate_limit_retry(
                notion.blocks.children.append,
                block_id=clean_id(page["id"]),
                children=chunk,
            )

        return page

    except Exception as e:
        return {"error": f"Failed to create Notion page: {str(e)}"}


# === Database Status/Priority Options ===

def get_database_status_priority(database_id, unified_token):
    """
    Fetch all status and priority-type (select) property options for a Notion database.
    Auto-detects properties by type.
    """
    notion = get_notion_client(unified_token)
    db_info = _with_notion_rate_limit_retry(
        notion.databases.retrieve, database_id=database_id
    )

    result = {
        "status_properties": {},
        "priority_properties": {}
    }

    for prop_name, prop_details in db_info["properties"].items():
        prop_type = prop_details.get("type")

        # Status fields
        if prop_type == "status":
            result["status_properties"][prop_name] = [
                opt["name"] for opt in prop_details["status"]["options"]
            ]

        # Priority-like fields (select)
        elif prop_type == "select":
            result["priority_properties"][prop_name] = [
                opt["name"] for opt in prop_details["select"]["options"]
            ]

    return result


def notion_get_database_schema(
    database_id: str,
    unified_token: str = None,
    sample_page_limit: int = 10,
) -> Dict[str, Any]:
    """
    Return a structured view of a database's schema:
    - All properties and their types
    - Select/status options (where applicable)
    - Example values observed from up to `sample_page_limit` rows
    """
    notion = get_notion_client(unified_token)

    # Clamp sample_page_limit to a safe range
    try:
        sample_page_limit = int(sample_page_limit)
    except (TypeError, ValueError):
        sample_page_limit = 10
    sample_page_limit = max(1, min(sample_page_limit, 100))

    try:
        db = _with_notion_rate_limit_retry(
            notion.databases.retrieve, database_id=database_id
        )
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to retrieve database schema: {str(e)}",
        }

    properties = db.get("properties", {}) or {}

    # Compute database title
    title_parts = db.get("title", []) or []
    pieces: list[str] = []
    for part in title_parts:
        if "plain_text" in part:
            pieces.append(part.get("plain_text") or "")
        elif "text" in part:
            pieces.append(part["text"].get("content", ""))
    db_title = "".join(pieces).strip() or None

    schema: Dict[str, Any] = {
        "success": True,
        "database_id": database_id,
        "title": db_title,
        "url": db.get("url"),
        "properties": {},
    }

    # Initialize per-property schema entries and example tracking
    examples_tracker: Dict[str, Dict[str, Any]] = {}

    for prop_name, prop_details in properties.items():
        prop_type = prop_details.get("type")
        entry: Dict[str, Any] = {
            "type": prop_type,
            "options": None,
            "examples": [],
        }

        if prop_type == "select":
            entry["options"] = [
                opt.get("name")
                for opt in prop_details.get("select", {}).get("options", [])
                if isinstance(opt, dict)
            ]
        elif prop_type == "multi_select":
            entry["options"] = [
                opt.get("name")
                for opt in prop_details.get("multi_select", {}).get("options", [])
                if isinstance(opt, dict)
            ]
        elif prop_type == "status":
            entry["options"] = [
                opt.get("name")
                for opt in prop_details.get("status", {}).get("options", [])
                if isinstance(opt, dict)
            ]

        schema["properties"][prop_name] = entry
        examples_tracker[prop_name] = {"seen": set(), "values": []}

    # Collect example values from up to `sample_page_limit` rows
    pages_seen = 0
    start_cursor = None

    while pages_seen < sample_page_limit:
        page_size = min(100, sample_page_limit - pages_seen)
        try:
            response = _with_notion_rate_limit_retry(
                notion.databases.query,
                database_id=database_id,
                start_cursor=start_cursor,
                page_size=page_size,
            )
        except Exception as e:
            # Don't fail the whole schema call because examples couldn't be collected
            schema["example_collection_error"] = str(e)
            break

        results = response.get("results", [])
        if not results:
            break

        for page in results:
            pages_seen += 1
            page_props = page.get("properties", {}) or {}

            for prop_name, prop_obj in page_props.items():
                if prop_name not in examples_tracker:
                    continue
                tracker = examples_tracker[prop_name]
                value = extract_property_value(prop_obj)
                if value is None:
                    continue

                # Normalize to a list so we can collect multiple scalar values from multi-valued props
                values_to_consider = value if isinstance(value, list) else [value]
                for v in values_to_consider:
                    # Use a JSON string token for deduplication but keep the original value
                    try:
                        token = json.dumps(v, sort_keys=True, default=str)
                    except TypeError:
                        token = str(v)

                    if token in tracker["seen"]:
                        continue
                    tracker["seen"].add(token)
                    tracker["values"].append(v)

                    # Cap examples per property to keep payload small
                    if len(tracker["values"]) >= 5:
                        break

        if not response.get("has_more"):
            break
        start_cursor = response.get("next_cursor")

    # Attach examples back onto the schema
    for prop_name, tracker in examples_tracker.items():
        values = tracker.get("values") or []
        schema["properties"][prop_name]["examples"] = values

    schema["sampled_pages"] = pages_seen
    return schema


# === Enhanced HTML to Notion with Rich Text Support ===

def clean_id_edit(id_str: str) -> str:
    """Ensure valid UUID format for Notion block IDs"""
    return str(uuid.UUID(re.sub(r"[^a-fA-F0-9]", "", id_str)))


def html_to_notion_blocks_edit(html: str) -> List[Dict]:
    """Convert HTML string to Notion blocks with styling + list structure."""
    if len(html) > 100_000:  # safety trim for huge HTML
        html = html[:100_000]

    soup = BeautifulSoup(html, "html.parser")
    blocks = []

    elements = soup.body.contents if soup.body else soup.contents

    for elem in elements:
        if isinstance(elem, str):
            continue
        blocks.extend(handle_element_edit(elem))

    return blocks


def _split_inline_bullets(text: str) -> tuple[str | None, List[str]]:
    """
    Heuristic: detect multiple "- item" style bullets that are written inline
    on a single line (e.g. "Next steps: - Do X - Do Y - Do Z") and split them
    into a prefix paragraph + individual bullet items.

    We only trigger this when there are at least two " - " separators, to avoid
    interfering with normal hyphen usage.
    """
    if not text:
        return None, []

    # Normalize internal whitespace for more reliable splitting
    compact = " ".join(text.split())
    parts = compact.split(" - ")

    # Need at least two bullets → three segments: prefix + bullet1 + bullet2...
    if len(parts) < 3:
        return None, []

    prefix_raw = parts[0].strip()
    bullets_raw = [p.strip() for p in parts[1:] if p.strip()]

    if not bullets_raw:
        return None, []

    prefix = prefix_raw or None
    return prefix, bullets_raw


def handle_element_edit(elem) -> List[Dict]:
    """Recursively process an HTML element into Notion blocks."""
    tag = elem.name
    if not tag:
        return []

    if tag in ["p", "span", "div"]:
        # Try to detect inline hyphen bullets that were not marked up as <ul>/<li>
        raw_text = elem.get_text(" ", strip=True)
        prefix, bullets = _split_inline_bullets(raw_text)

        blocks: List[Dict] = []

        # Add the non-list prefix (e.g. "Next steps:") as a normal paragraph
        if prefix:
            blocks.append(
                paragraph_block_edit(
                    [{"type": "text", "text": {"content": prefix[:2000]}}]
                )
            )

        # Convert detected bullets into proper Notion bulleted_list_item blocks
        for btxt in bullets:
            blocks.append(
                {
                    "object": "block",
                    "type": "bulleted_list_item",
                    "bulleted_list_item": {
                        "rich_text": [
                            {
                                "type": "text",
                                "text": {"content": btxt[:2000]},
                            }
                        ],
                        "children": [],
                    },
                }
            )

        # Fallback: no bullets detected → treat as a normal paragraph preserving rich text
        if not blocks:
            text_fragments = extract_rich_text_edit(elem)
            return [paragraph_block_edit(text_fragments)]

        return blocks

    elif tag in ["h1", "h2", "h3"]:
        text_fragments = extract_rich_text_edit(elem)
        return [heading_block_edit(text_fragments, int(tag[1]))]

    elif tag == "ul":
        return [list_item_block_edit(li, "bulleted_list_item") for li in elem.find_all("li", recursive=False)]

    elif tag == "ol":
        return [list_item_block_edit(li, "numbered_list_item") for li in elem.find_all("li", recursive=False)]

    elif tag == "img":
        return [image_block_edit(elem)]

    elif tag in ["pre", "code"]:
        return [code_block_edit(elem.get_text())]

    elif tag == "table":
        return [table_block_edit(elem)]

    else:
        return [paragraph_block_edit([{
            "type": "text",
            "text": {"content": f"[UNSUPPORTED <{tag}>]"}
        }])]


def extract_rich_text_edit(elem) -> List[Dict]:
    """Extract rich text with annotations (bold, italic, underline)."""
    fragments = []
    for node in elem.descendants:
        if node.name is None:  # text node
            text = node.strip()
            if not text:
                continue

            parent_tags = [p.name for p in node.parents if p.name]
            annotations = {
                "bold": "b" in parent_tags or "strong" in parent_tags,
                "italic": "i" in parent_tags or "em" in parent_tags,
                "underline": "u" in parent_tags,
                "strikethrough": "s" in parent_tags or "del" in parent_tags,
                "code": "code" in parent_tags
            }

            fragments.append({
                "type": "text",
                "text": {"content": text[:2000]},
                "annotations": annotations
            })

    return fragments or [{"type": "text", "text": {"content": ""}}]


def paragraph_block_edit(rich_text: List[Dict]) -> Dict:
    return {
        "object": "block",
        "type": "paragraph",
        "paragraph": {"rich_text": rich_text}
    }


def heading_block_edit(rich_text: List[Dict], level: int) -> Dict:
    heading_type = f"heading_{min(level, 3)}"
    return {
        "object": "block",
        "type": heading_type,
        heading_type: {"rich_text": rich_text}
    }


def list_item_block_edit(li_tag, list_type: str) -> Dict:
    text_fragments = extract_rich_text_edit(li_tag)

    # Handle nested lists inside this li
    children_blocks = []
    for child in li_tag.find_all(["ul", "ol"], recursive=False):
        children_blocks.extend(handle_element_edit(child))

    return {
        "object": "block",
        "type": list_type,
        list_type: {
            "rich_text": text_fragments,
            "children": children_blocks if children_blocks else []
        },
    }


def image_block_edit(img_tag) -> Dict:
    src = img_tag.get("src", "").strip()
    if not src:
        return paragraph_block_edit([{"type": "text", "text": {"content": "[MISSING IMAGE SRC]"}}])

    return {
        "object": "block",
        "type": "image",
        "image": {
            "type": "external",
            "external": {"url": src}
        },
    }


def code_block_edit(code_text: str) -> Dict:
    return {
        "object": "block",
        "type": "code",
        "code": {
            "rich_text": [{"type": "text", "text": {"content": code_text[:2000]}}],
            "language": "plain text"
        },
    }


def table_block_edit(table_tag) -> Dict:
    """Represent table as code block (for now)."""
    rows = []
    for row_tag in table_tag.find_all("tr", recursive=False):
        cells = row_tag.find_all(["td", "th"], recursive=False)
        cell_texts = [cell.get_text(strip=True)[:2000] for cell in cells]
        if cell_texts:
            rows.append(" | ".join(cell_texts))

    table_string = "\n".join(rows) if rows else "[EMPTY TABLE]"
    return code_block_edit(table_string)


# === Page Update Function ===

def update_notion_page(unified_token: str, page_id: str, html_string: str, title: str = None) -> Dict:
    notion = get_notion_client(unified_token)
    new_blocks = html_to_notion_blocks_edit(html_string)

    try:
        # 1. Update title if provided
        if title:
            _with_notion_rate_limit_retry(
                notion.pages.update,
                page_id=clean_id_edit(page_id),
                properties={
                    "title": {
                        "title": [
                            {"type": "text", "text": {"content": title[:2000]}}
                        ]
                    }
                },
            )

        # 2. Fetch existing blocks
        existing_blocks = _with_notion_rate_limit_retry(
            notion.blocks.children.list,
            block_id=clean_id_edit(page_id),
        ).get("results", [])
        min_len = min(len(existing_blocks), len(new_blocks))

        # 3. Update where possible
        for i in range(min_len):
            old_block = existing_blocks[i]
            new_block = new_blocks[i]

            old_type = old_block.get("type")
            new_type = new_block.get("type")

            if old_type and new_type and old_type == new_type:
                # Same block type → update this block in place using the type-specific payload
                block_payload = new_block.get(new_type, {})
                _with_notion_rate_limit_retry(
                    notion.blocks.update,
                    block_id=clean_id_edit(old_block["id"]),
                    **{new_type: block_payload},
                )
            else:
                # Different type → delete old, insert new
                _with_notion_rate_limit_retry(
                    notion.blocks.delete, block_id=clean_id_edit(old_block["id"])
                )
                _with_notion_rate_limit_retry(
                    notion.blocks.children.append,
                    block_id=clean_id_edit(page_id),
                    children=[new_block],
                )

        # 4. Append extra blocks
        if len(new_blocks) > len(existing_blocks):
            extra_blocks = new_blocks[len(existing_blocks):]
            for i in range(0, len(extra_blocks), MAX_BLOCKS_PER_REQUEST):
                _with_notion_rate_limit_retry(
                    notion.blocks.children.append,
                    block_id=clean_id_edit(page_id),
                    children=extra_blocks[i:i + MAX_BLOCKS_PER_REQUEST],
                )

        # 5. Delete extra old blocks
        elif len(existing_blocks) > len(new_blocks):
            for i in range(len(new_blocks), len(existing_blocks)):
                _with_notion_rate_limit_retry(
                    notion.blocks.delete,
                    block_id=clean_id_edit(existing_blocks[i]["id"]),
                )

        return {"success": True, "message": "Page updated successfully", "page_id": page_id}

    except Exception as e:
        return {"error": f"Failed to update Notion page: {str(e)}"}


# === Document Deletion Function ===

def notion_delete_document(page_id, unified_token):
    """
    Soft-delete a Notion page by archiving it.
    """
    notion = get_notion_client(unified_token)
    try:
        response = _with_notion_rate_limit_retry(
            notion.pages.update,
            page_id=page_id,
            archived=True,
        )
        return {"success": True, "page_id": page_id, "archived": True}
    except Exception as e:
        return {"success": False, "page_id": page_id, "error": str(e)}


def notion_create_database(
    parent_id: str,
    title: str,
    unified_token: str = None,
    properties: Optional[Dict[str, Any]] = None,
) -> Dict:
    """
    Create a new Notion database under any Notion page.

    Notes:
    - The parent can be any page_id, including:
      - Top-level workspace pages
      - Child pages under other pages
      - Pages that are rows inside an existing database (database → page → database ... onion structure)
    - This mirrors Notion's ability to nest databases inside pages at any depth.
    """
    notion = get_notion_client(unified_token)
    
    try:
        # Verify parent is a page
        _with_notion_rate_limit_retry(notion.pages.retrieve, parent_id)
        
        # If caller did not provide a schema, create a minimal one: only
        # a title property. All other fields should be defined explicitly
        # by the caller via the properties argument to keep schemas
        # aligned with the requested topic (no generic Status/Priority).
        db_properties: Dict[str, Any] = properties or {
            "Name": {
                "title": {},
            }
        }

        database = _with_notion_rate_limit_retry(
            notion.databases.create,
            parent={"type": "page_id", "page_id": parent_id},
            title=[
                {
                    "type": "text",
                    "text": {"content": title}
                }
            ],
            properties=db_properties,
        )
        return {"success": True, "id": database["id"], "url": database["url"], "type": "database"}
    except Exception as e:
        return {"success": False, "error": f"Failed to create database: {str(e)}"}


def text_to_notion_blocks(text: str) -> List[Dict]:
    """
    Parse plain text into Notion blocks with proper formatting.
    Supports:
    - Headings: # H1, ## H2, ### H3
    - Bullet lists: - or * 
    - Numbered lists: 1. or 1)
    - Dividers: --- or ***
    - Paragraphs: regular text
    """
    if not text:
        return []
    
    blocks = []
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (they'll be handled as paragraph breaks)
        if not line:
            i += 1
            continue
        
        # Check for dividers
        if line in ['---', '***', '___'] or line.replace('-', '').replace('*', '').replace('_', '') == '':
            blocks.append({
                "object": "block",
                "type": "divider",
                "divider": {}
            })
            i += 1
            continue
        
        # Check for headings
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            if level <= 3 and level > 0:
                heading_text = line[level:].strip()
                if heading_text:
                    blocks.append(heading_block(heading_text, level))
                    i += 1
                    continue
        
        # Check for bullet lists
        if line.startswith('- ') or line.startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item_text = lines[i].strip()[2:].strip()  # Remove '- ' or '* '
                if item_text:
                    list_items.append({
                        "object": "block",
                        "type": "bulleted_list_item",
                        "bulleted_list_item": {
                            "rich_text": [{"type": "text", "text": {"content": item_text[:2000]}}]
                        }
                    })
                i += 1
            blocks.extend(list_items)
            continue
        
        # Check for numbered lists
        numbered_match = re.match(r'^(\d+)[.)]\s+(.+)$', line)
        if numbered_match:
            list_items = []
            expected_num = 1
            while i < len(lines):
                match = re.match(r'^(\d+)[.)]\s+(.+)$', lines[i].strip())
                if match:
                    num = int(match.group(1))
                    if num == expected_num:
                        item_text = match.group(2).strip()
                        if item_text:
                            list_items.append({
                                "object": "block",
                                "type": "numbered_list_item",
                                "numbered_list_item": {
                                    "rich_text": [{"type": "text", "text": {"content": item_text[:2000]}}]
                                }
                            })
                        expected_num += 1
                        i += 1
                    else:
                        break
                else:
                    break
            blocks.extend(list_items)
            continue
        
        # Regular paragraph - collect consecutive non-empty lines
        paragraph_lines = []
        while i < len(lines) and lines[i].strip():
            line_stripped = lines[i].strip()
            # Don't include if it's a special format
            if (line_stripped.startswith('#') or 
                line_stripped.startswith('- ') or 
                line_stripped.startswith('* ') or
                re.match(r'^\d+[.)]\s+', line_stripped) or
                line_stripped in ['---', '***', '___']):
                break
            paragraph_lines.append(lines[i].rstrip())
            i += 1
        
        if paragraph_lines:
            paragraph_text = '\n'.join(paragraph_lines).strip()
            if paragraph_text:
                blocks.append(paragraph_block(paragraph_text))
            continue
        
        i += 1
    
    return blocks


def notion_create_parent_page(title: str, parent_page_id: str = None, content: str = "", unified_token: str = None) -> Dict:
    """Create a new parent page. If parent_page_id is provided, creates as child page. Otherwise tries workspace-level creation."""
    notion = get_notion_client(unified_token)
    
    try:
        # Determine parent type
        if parent_page_id:
            # Create as child page of existing page
            parent = {"type": "page_id", "page_id": parent_page_id}
        else:
            # Try workspace-level creation (may fail for internal integrations)
            parent = {"type": "workspace", "workspace": True}
        
        page = _with_notion_rate_limit_retry(
            notion.pages.create,
            parent=parent,
            properties={"title": {"title": [{"text": {"content": title}}]}},
            children=[
                {
                    "object": "block",
                    "type": "paragraph",
                    "paragraph": {
                        "rich_text": [{"type": "text", "text": {"content": content}}]
                    },
                }
            ] if content else [],
        )
        return {
            "success": True, 
            "id": page["id"], 
            "url": page["url"], 
            "type": "parent_page",
            "title": title,
            "parent_type": "workspace" if not parent_page_id else "child_page"
        }
    except Exception as e:
        error_msg = str(e)
        if "workspace-level private pages is not supported" in error_msg:
            return {
                "success": False, 
                "error": "Cannot create workspace-level pages with internal integrations. Please provide a parent_page_id to create as child page, or use a public integration with insert_content capability.",
                "suggestion": "Use list_parent_pages() to find an existing page to use as parent, or create the page manually in Notion first."
            }
        return {"success": False, "error": f"Failed to create parent page: {error_msg}"}


def notion_list_parent_pages(unified_token: str = None) -> List[Dict]:
    """List top-level parent pages that can contain databases (not database rows)"""
    notion = get_notion_client(unified_token)
    
    try:
        # Search for all pages
        results = notion.search(filter={"property": "object", "value": "page"})["results"]
        parent_pages = []
        
        for page in results:
            # Check if this is a top-level page (not a database row)
            parent = page.get("parent", {})
            parent_type = parent.get("type")
            
            # Only include pages that are NOT in databases (i.e., top-level pages or child pages)
            if parent_type in ["workspace", "page_id"]:
                title = get_title_from_properties(page.get("properties", {}))
                parent_pages.append({
                    "id": page["id"],
                    "title": title,
                    "url": page.get("url"),
                    "parent_type": parent_type,
                    "last_edited": page.get("last_edited_time")
                })
        
        # Sort by last edited time (most recent first)
        parent_pages.sort(key=lambda x: x.get("last_edited", ""), reverse=True)
        return parent_pages
        
    except Exception as e:
        print(f"Error listing parent pages: {str(e)}")
        return []


def notion_list_all_database_parents(unified_token: str = None) -> List[Dict]:
    """
    List pages that can serve as parents for a database creation, including:
    - Top-level workspace pages
    - Child pages under other pages
    - Pages that are rows inside databases (parent.type == "database_id")

    Returns a simplified list: [{id, title, url, parent_type, last_edited}]
    """
    notion = get_notion_client(unified_token)
    try:
        results = notion.search(filter={"property": "object", "value": "page"}).get("results", [])
        pages: List[Dict] = []
        for page in results:
            parent = page.get("parent", {})
            parent_type = parent.get("type")
            # Accept all parent types that are pages (workspace, page_id) and also pages in databases (database_id)
            if parent_type in ["workspace", "page_id", "database_id"]:
                title = get_title_from_properties(page.get("properties", {}))
                pages.append({
                    "id": page.get("id"),
                    "title": title,
                    "url": page.get("url"),
                    "parent_type": parent_type,
                    "last_edited": page.get("last_edited_time")
                })
        pages.sort(key=lambda x: x.get("last_edited", ""), reverse=True)
        return pages
    except Exception as e:
        print(f"Error listing all potential database parents: {str(e)}")
        return []


def notion_list_child_pages(page_id: str, unified_token: str = None) -> List[Dict]:
    """List child pages and databases within a parent page"""
    notion = get_notion_client(unified_token)
    
    try:
        # Get all child blocks of the page
        children = []
        has_more = True
        start_cursor = None
        
        while has_more:
            response = notion.blocks.children.list(
                block_id=page_id,
                start_cursor=start_cursor
            )
            children.extend(response.get("results", []))
            has_more = response.get("has_more", False)
            start_cursor = response.get("next_cursor")
        
        # Filter for child pages and databases
        child_items = []
        for child in children:
            if child.get("type") == "child_page":
                child_items.append({
                    "id": child["id"],
                    "title": child["child_page"]["title"],
                    "type": "page"
                })
            elif child.get("type") == "child_database":
                child_items.append({
                    "id": child["id"],
                    "title": child["child_database"]["title"],
                    "type": "database"
                })
        
        return child_items
    except Exception as e:
        return []
